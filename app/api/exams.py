"""
Exam management API endpoints.
"""
from datetime import datetime, timezone, timedelta
from decimal import Decimal
from typing import List, Optional, Dict, Any, Mapping, Set, Tuple
import asyncio
import csv
import html
import io
import logging
import random
import hashlib  # Added for stable seeding
import re
import time
from types import SimpleNamespace
import sqlalchemy
from fastapi import APIRouter, Depends, HTTPException, status, Request, Query
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete, func, or_, and_, text, update
from sqlalchemy.orm import selectinload, joinedload, noload
from pydantic import BaseModel
import pytz
import uuid

from app.database import get_db, get_db_read
from app.models.user import User
from app.models.exam import Exam
from app.models.question import Question, QuestionOption
from app.models.session import ExamSession, Answer, ExamLog
from app.models.activity_log import UserActivityLog
from app.schemas.exam import (
    ExamCreate, ExamResponse,
    QuestionResponse, QuestionOptionResponse,
    ExamStartResponse, ExamAnalytics
)
from app.schemas.answer import ExamSubmitResponse
from app.core.security import (
    AuthenticatedUser,
    get_current_user,
    get_current_user_hot_path,
    get_current_teacher,
    get_current_exam_monitor,
    create_session_poll_token,
    is_pengawas_user,
)
from app.middleware.seb_validation import validate_seb_headers, get_client_info
from app.core.redis_pubsub import (
    publish_message, store_session_data, get_session_data,
    get_redis, update_session_activity
)
from app.core.client_ip import get_client_ip
from app.core.answer_review_helpers import (
    QUESTION_TYPE_LABELS,
    build_option_map as _build_option_map,
    coerce_bool as _coerce_bool,
    resolve_question_statements as _resolve_question_statements,
    resolve_statement_keys as _resolve_statement_keys,
    status_from_answer as _status_from_answer,
)
from app.core.exam_access_policy import (
    ensure_exam_participant_access as _ensure_exam_participant_access,
    is_exam_participant_role as _is_exam_participant_role,
    participant_has_exam_access as _participant_has_exam_access,
)
from app.core.exam_results_cache import (
    build_exam_results_cache_key as _build_exam_results_cache_key,
    build_exam_results_viewer_scope as _build_exam_results_viewer_scope,
    get_cached_exam_results as _get_cached_exam_results,
    invalidate_exam_results_cache as _invalidate_exam_results_cache,
    set_cached_exam_results as _set_cached_exam_results,
)
from app.core.exam_runtime_cache import (
    get_session_answer_count_cached as _get_session_answer_count_cached,
)
from app.core.monitoring_delta import publish_monitoring_delta
from app.core.exam_session_helpers import (
    calculate_effective_timer,
    merge_statement_answer_metadata,
    parse_iso_datetime_utc,
    resolve_timer_context,
    safe_int,
)
from app.core.roles import (
    ROLE_ADMIN,
    ROLE_DEVELOPER,
    is_developer_exam_hidden_for_viewer,
    is_developer_role,
    normalize_role,
)
from app.core.session_recovery import RECOVERY_CATEGORY_ADMIN, evaluate_session_recovery
from app.core.start_db_admission import bind_start_admission, start_db_segment
from app.config import settings
from app.core.feature_flags import require_feature_enabled
from app.core.rate_limiter import RateLimiters, check_rate_limit
from app.services.exam_service import ExamService
from app.services.exam_submission_service import finalize_exam_session_submission

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/exams", tags=["Exams"])
public_router = APIRouter(prefix="/api/exams", tags=["Exams"])

# Admin Audit Logging Helper
async def log_admin_action(
    db: AsyncSession,
    admin_user: User,
    action: str,
    target_type: str,
    target_id: int,
    target_name: str,
    details: Optional[Dict[str, Any]] = None
):

    """Log admin actions for audit trail."""
    if normalize_role(admin_user.role) in {ROLE_ADMIN, ROLE_DEVELOPER}:
        log_entry = UserActivityLog(
            user_id=admin_user.id,
            event_type=f"admin_{action}",
            event_data={
                "admin_username": admin_user.username,
                "action": action,
                "target_type": target_type,
                "target_id": target_id,
                "target_name": target_name,
                "details": details or {}
            }
        )
        db.add(log_entry)
        await db.commit()
        logger.warning(f"🚨 ADMIN ACTION: {admin_user.username} {action} {target_type} '{target_name}' (ID: {target_id})")


EXAM_CRITICAL_METADATA_FIELDS = (
    "start_time",
    "end_time",
    "allowed_classes",
    "allowed_students",
)

EXAM_AUDIT_FIELDS = (
    "title",
    "description",
    "duration_minutes",
    "start_time",
    "end_time",
    "passing_score",
    "max_attempts",
    "shuffle_questions",
    "shuffle_options",
    "show_results",
    "allow_review",
    "is_published",
    "subject",
    "exam_type",
    "academic_year",
    "show_teacher_name",
    "allowed_classes",
    "allowed_students",
)


def _normalize_csv_restriction_for_compare(value: Optional[str]) -> Tuple[str, ...]:
    if not value:
        return tuple()
    return tuple(sorted({part.strip().upper() for part in value.split(",") if part.strip()}))


def _datetime_for_compare(value: Optional[datetime]) -> Optional[datetime]:
    if value is None:
        return None
    if value.tzinfo is None:
        normalized = value.replace(tzinfo=timezone.utc)
    else:
        normalized = value.astimezone(timezone.utc)
    return normalized.replace(microsecond=0)


def _exam_datetime_changed(current_value: Optional[datetime], new_value: Optional[datetime]) -> bool:
    return _datetime_for_compare(current_value) != _datetime_for_compare(new_value)


def _exam_critical_metadata_changes(exam: Exam, exam_data: ExamCreate) -> List[str]:
    changed: List[str] = []
    if _exam_datetime_changed(exam.start_time, exam_data.start_time):
        changed.append("start_time")
    if _exam_datetime_changed(exam.end_time, exam_data.end_time):
        changed.append("end_time")
    if _normalize_csv_restriction_for_compare(exam.allowed_classes) != _normalize_csv_restriction_for_compare(exam_data.allowed_classes):
        changed.append("allowed_classes")
    if _normalize_csv_restriction_for_compare(exam.allowed_students) != _normalize_csv_restriction_for_compare(exam_data.allowed_students):
        changed.append("allowed_students")
    return changed


def _json_safe_audit_value(value: Any) -> Any:
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    if isinstance(value, (list, tuple, set)):
        return [_json_safe_audit_value(item) for item in value]
    if isinstance(value, dict):
        return {str(key): _json_safe_audit_value(item) for key, item in value.items()}
    return str(value)


def _collect_exam_update_changes(original_values: Dict[str, Any], exam_data: ExamCreate) -> Dict[str, Dict[str, Any]]:
    changes: Dict[str, Dict[str, Any]] = {}
    for field in EXAM_AUDIT_FIELDS:
        old_value = original_values.get(field)
        new_value = getattr(exam_data, field)
        if field in {"start_time", "end_time"}:
            changed = _exam_datetime_changed(old_value, new_value)
        elif field in {"allowed_classes", "allowed_students"}:
            changed = _normalize_csv_restriction_for_compare(old_value) != _normalize_csv_restriction_for_compare(new_value)
        else:
            changed = _json_safe_audit_value(old_value) != _json_safe_audit_value(new_value)
        if changed:
            changes[field] = {
                "old": _json_safe_audit_value(old_value),
                "new": _json_safe_audit_value(new_value),
            }
    return changes


def _student_summary_entry(row: Any) -> Dict[str, Any]:
    return {
        "id": int(row["id"]),
        "username": row["username"] or "",
        "full_name": row["full_name"] or row["username"] or "",
        "student_class": row["student_class"] or "",
    }


def _format_wib_datetime(value: Optional[datetime]) -> str:
    if value is None:
        return ""
    wib = pytz.timezone("Asia/Jakarta")
    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone.utc)
    return value.astimezone(wib).strftime("%Y-%m-%d %H:%M:%S WIB")

def get_exam_service(db: AsyncSession = Depends(get_db)) -> ExamService:
    return ExamService(db)

def get_exam_service_read(db: AsyncSession = Depends(get_db_read)) -> ExamService:
    """Service using Read Replica"""
    return ExamService(db)


async def _get_exam_creator_role(db: AsyncSession, creator_id: Optional[int]) -> Optional[str]:
    if not creator_id:
        return None

    creator_role_result = await db.execute(
        select(User.role).where(User.id == creator_id)
    )
    return creator_role_result.scalar_one_or_none()


def _raise_hidden_exam_error() -> None:
    raise HTTPException(status_code=404, detail="Ujian tidak ditemukan")


def _enforce_developer_exam_visibility(current_user: User, exam_creator_role: Optional[str]) -> None:
    if is_developer_exam_hidden_for_viewer(current_user.role, exam_creator_role):
        _raise_hidden_exam_error()


async def _enforce_exam_owner_or_admin_access(
    db: AsyncSession,
    current_user: User,
    exam_creator_id: int,
    *,
    allow_pengawas: bool = False,
) -> str:
    creator_role = await _get_exam_creator_role(db, exam_creator_id)
    _enforce_developer_exam_visibility(current_user, creator_role)

    if exam_creator_id == current_user.id:
        return str(creator_role or "")

    if bool(getattr(current_user, "is_admin", False)):
        return str(creator_role or "")

    if allow_pengawas and is_pengawas_user(current_user):
        return str(creator_role or "")

    raise HTTPException(status_code=403, detail="Tidak memiliki akses")


def _pick_latest_scored_exam_session_per_user(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Pick one session row per user for exam result views.

    Selection rule:
    1. Prefer latest row with non-null score for each user.
    2. Fallback to the user's latest row when all scores are null.
    """
    latest_any: Dict[int, Dict[str, Any]] = {}
    latest_scored: Dict[int, Dict[str, Any]] = {}

    for row in rows:
        user_id = int(row["user_id"])
        if user_id not in latest_any:
            latest_any[user_id] = row
        if row.get("score") is not None and user_id not in latest_scored:
            latest_scored[user_id] = row

    selected_rows: List[Dict[str, Any]] = []
    for user_id, any_row in latest_any.items():
        selected_rows.append(latest_scored.get(user_id, any_row))

    min_datetime_utc = datetime.min.replace(tzinfo=timezone.utc)
    selected_rows.sort(
        key=lambda row: (
            row.get("end_time") is not None,
            row.get("end_time") or min_datetime_utc,
            int(row.get("session_id") or 0),
        ),
        reverse=True,
    )
    return selected_rows


EXAM_START_VALIDATION_CACHE_PREFIX = "cache:exam-start-validation:v1"
EXAM_START_VALIDATION_CACHE_TTL_SECONDS = 120
EXAM_START_VALIDATION_LOCAL_CACHE_TTL_SECONDS = 300
SESSION_POLL_TOKEN_EXPIRES_MINUTES = 15
SESSION_WRITE_LOCK_NAMESPACE = 48102

_exam_start_validation_local_cache: Dict[int, float] = {}


def _build_exam_start_validation_cache_key(exam_id: int) -> str:
    """Build the Redis cache key for exam-start option integrity validation."""
    return f"{EXAM_START_VALIDATION_CACHE_PREFIX}:{int(exam_id)}"


async def _ensure_exam_start_option_integrity(db: AsyncSession, exam_id: int) -> None:
    """
    Ensure option-based questions are renderable before exam start.
    Uses short Redis cache to avoid repeating heavy validation on burst starts.
    """
    now_monotonic = time.monotonic()
    local_cached_until = _exam_start_validation_local_cache.get(exam_id, 0.0)
    if now_monotonic < local_cached_until:
        return

    cache_key = _build_exam_start_validation_cache_key(exam_id)
    lock_key = f"{cache_key}:lock"
    redis = None
    validation_lock_token: Optional[str] = None
    try:
        redis = await get_redis()
        cached = await redis.get(cache_key)
        if cached == "1":
            _exam_start_validation_local_cache[exam_id] = (
                now_monotonic + EXAM_START_VALIDATION_LOCAL_CACHE_TTL_SECONDS
            )
            return

        validation_lock_token = uuid.uuid4().hex
        acquired = await redis.set(lock_key, validation_lock_token, ex=15, nx=True)
        if not acquired:
            wait_deadline = time.monotonic() + 12.0
            while time.monotonic() < wait_deadline:
                cached = await redis.get(cache_key)
                if cached == "1":
                    _exam_start_validation_local_cache[exam_id] = (
                        time.monotonic() + EXAM_START_VALIDATION_LOCAL_CACHE_TTL_SECONDS
                    )
                    return
                await asyncio.sleep(0.05)

            logger.warning(
                "Exam start validation lock wait expired for exam %s; skip duplicate validation on hot path",
                exam_id,
            )
            return
    except Exception as exc:
        logger.warning(
            "Failed reading exam start validation cache for exam %s: %s",
            exam_id,
            str(exc),
        )

    async with start_db_segment("integrity"):
        orphaned_check = await db.execute(
            select(Question.id, Question.question_text, Question.question_type)
            .outerjoin(QuestionOption, Question.id == QuestionOption.question_id)
            .where(
                Question.exam_id == exam_id,
                QuestionOption.id == None,
                or_(
                    Question.question_type.in_(["multiple_choice", "true_false"]),
                    and_(
                        Question.question_type == "multiple_choice_complex",
                        func.coalesce(Question.pgk_type, "checkbox") != "table_validation",
                    ),
                ),
            )
            .group_by(Question.id, Question.question_text, Question.question_type)
        )
        orphaned_questions = orphaned_check.all()
    if orphaned_questions:
        orphaned_ids = [str(q[0]) for q in orphaned_questions]
        logger.error(
            "EXAM_START | INVALID EXAM %s | Questions with 0 options: %s",
            exam_id,
            orphaned_ids,
        )
        raise HTTPException(
            status_code=500,
            detail=(
                f"Ujian memiliki {len(orphaned_questions)} soal pilihan ganda tanpa pilihan jawaban. "
                f"Tidak bisa dimulai. Silakan hubungi pengawas atau administrator."
            ),
        )

    _exam_start_validation_local_cache[exam_id] = (
        time.monotonic() + EXAM_START_VALIDATION_LOCAL_CACHE_TTL_SECONDS
    )

    if redis is not None:
        try:
            await redis.set(cache_key, "1", ex=EXAM_START_VALIDATION_CACHE_TTL_SECONDS)
        except Exception as exc:
            logger.warning(
                "Failed writing exam start validation cache for exam %s: %s",
                exam_id,
                str(exc),
            )
        finally:
            if validation_lock_token is not None:
                try:
                    current_lock = await redis.get(lock_key)
                    if current_lock == validation_lock_token:
                        await redis.delete(lock_key)
                except Exception as exc:
                    logger.warning(
                        "Failed releasing exam start validation lock for exam %s: %s",
                        exam_id,
                        str(exc),
                    )


def _is_placeholder_question(settings_payload: Optional[Dict[str, Any]]) -> bool:
    settings_dict = settings_payload or {}
    return bool(settings_dict.get("is_placeholder", False))


def _can_shuffle_placeholder_options(
    settings_payload: Optional[Dict[str, Any]],
    *,
    has_image: bool = False
) -> bool:
    """Allow placeholder shuffle only for non-image sources."""
    settings_dict = settings_payload or {}
    if not _is_placeholder_question(settings_dict):
        return False

    placeholder_source = str(settings_dict.get("placeholder_source") or "").strip().lower()
    allow_flag = bool(settings_dict.get("allow_placeholder_shuffle", False))
    if has_image or placeholder_source == "image":
        # Mode 2 is disabled globally; image-based placeholders stay fixed.
        return False

    return allow_flag


def _stable_shuffle_with_seed(items: List[Any], seed_str: str) -> List[Any]:
    """
    Deterministic shuffle with a guaranteed order change when possible.
    If RNG returns original order, rotate deterministically to enforce variation.
    """
    original_items = list(items)
    if len(original_items) < 2:
        return original_items

    seed = int(hashlib.md5(seed_str.encode()).hexdigest(), 16)
    rng = random.Random(seed)
    shuffled_items = list(original_items)
    rng.shuffle(shuffled_items)

    if shuffled_items == original_items:
        offset = (seed % (len(shuffled_items) - 1)) + 1
        shuffled_items = shuffled_items[offset:] + shuffled_items[:offset]

    return shuffled_items


async def _autofill_placeholder_options_for_publish(exam_id: int, db: AsyncSession) -> int:
    """
    Auto-fill A/B/C/... placeholders for option-based questions when teacher only sets answer keys.
    This keeps student rendering/scoring safe while allowing very fast question authoring.
    """
    result = await db.execute(
        select(Question)
        .options(selectinload(Question.options))
        .where(Question.exam_id == exam_id)
        .order_by(Question.order_index.asc(), Question.id.asc())
    )
    questions = result.scalars().all()

    def _has_embedded_choice_lines(raw_text: Optional[str]) -> bool:
        text = raw_text or ""
        text = re.sub(r"(?i)<br\s*/?>", "\n", text)
        text = re.sub(r"(?i)</(p|div|li)>", "\n", text)
        text = re.sub(r"<[^>]+>", " ", text)
        labels = {
            match.group(1).upper()
            for match in re.finditer(r"(?i)\b([A-Za-z])[.):]\s+\S+", text)
        }
        return len(labels) >= 2

    MC_MIN_OPTIONS = 3
    PGK_CHECKBOX_MIN_OPTIONS = 4
    updated_questions = 0

    for q in questions:
        settings = dict(q.question_settings or {})
        pgk_type = q.pgk_type or settings.get("pgk_type", "checkbox")
        minimum_required = (
            MC_MIN_OPTIONS
            if q.question_type == "multiple_choice"
            else PGK_CHECKBOX_MIN_OPTIONS
        )

        is_option_based = (
            q.question_type == "multiple_choice" or
            (q.question_type == "multiple_choice_complex" and pgk_type != "table_validation")
        )
        if not is_option_based or not q.options or len(q.options) < minimum_required:
            continue

        sorted_opts = sorted(q.options, key=lambda o: o.order_index)
        real_count = sum(1 for opt in sorted_opts if (opt.option_text or "").strip())
        if real_count >= minimum_required:
            continue

        has_selected_key = any(bool(opt.is_correct) for opt in sorted_opts)
        has_embedded = _has_embedded_choice_lines(q.question_text)
        has_media = bool(q.image_url)

        # Super-permissive mode: if key is selected, normalize blank options into A/B/C...
        if not has_selected_key and not has_embedded and not has_media:
            continue

        changed = False
        for idx, opt in enumerate(sorted_opts):
            if (opt.option_text or "").strip():
                continue
            letter = chr(65 + (idx % 26))
            suffix = str(idx // 26 + 1) if idx >= 26 else ""
            opt.option_text = f"{letter}{suffix}"
            changed = True

        if changed:
            settings["is_placeholder"] = True
            if has_media:
                settings["placeholder_source"] = "image"
            elif has_embedded:
                settings["placeholder_source"] = "question_text"
            else:
                settings["placeholder_source"] = "auto_no_option"
            q.question_settings = settings
            updated_questions += 1

    if updated_questions > 0:
        await db.flush()

    return updated_questions


async def _validate_questions_for_publish(exam_id: int, db: AsyncSession) -> None:
    """Validate question completeness before exam publish."""
    result = await db.execute(
        select(Question)
        .options(selectinload(Question.options))
        .where(Question.exam_id == exam_id)
        .order_by(Question.order_index.asc(), Question.id.asc())
    )
    questions = result.scalars().all()

    if not questions:
        raise HTTPException(status_code=400, detail="Ujian belum memiliki soal")

    errors: List[str] = []

    MC_MIN_OPTIONS = 3
    PGK_CHECKBOX_MIN_OPTIONS = 4

    def _is_generated_placeholder(option_text: Optional[str]) -> bool:
        normalized = (option_text or "").strip().upper()
        if not normalized:
            return False
        return bool(re.fullmatch(r"[A-Z](?:[2-9][0-9]*)?", normalized))

    def _count_real_options(options: List[QuestionOption]) -> int:
        count = 0
        for opt in options:
            text = (opt.option_text or "").strip()
            if not text:
                continue
            if _is_generated_placeholder(text):
                continue
            count += 1
        return count

    def _has_embedded_choice_lines(raw_text: Optional[str]) -> bool:
        """Detect copy-pasted options inside question text (e.g. A. ..., B. ...)."""
        text = raw_text or ""
        # Keep line boundaries when source text contains simple HTML tags.
        text = re.sub(r"(?i)<br\s*/?>", "\n", text)
        text = re.sub(r"(?i)</(p|div|li)>", "\n", text)
        text = re.sub(r"<[^>]+>", " ", text)
        labels = {
            match.group(1).upper()
            for match in re.finditer(r"(?i)\b([A-Za-z])[.):]\s+\S+", text)
        }
        return len(labels) >= 2

    for idx, q in enumerate(questions, 1):
        q_type = q.question_type
        settings = q.question_settings or {}
        pgk_type = q.pgk_type or settings.get("pgk_type", "checkbox")

        question_text = (q.question_text or "").strip()
        has_media = bool(q.image_url or q.video_url or q.audio_url)
        if not question_text and not has_media:
            errors.append(f"Soal No. {idx}: Pertanyaan masih kosong")

        if q_type == "multiple_choice":
            real_options_count = _count_real_options(q.options)
            has_embedded_options = _has_embedded_choice_lines(q.question_text)
            has_correct = any(opt.is_correct for opt in q.options)
            is_image_mode = bool(q.image_url)
            permissive_key_only_mode = has_correct and not is_image_mode and not has_embedded_options

            if real_options_count < MC_MIN_OPTIONS and not is_image_mode and not has_embedded_options and not permissive_key_only_mode:
                errors.append(
                    f"Soal No. {idx} (Pilihan Ganda): Minimal harus ada {MC_MIN_OPTIONS} opsi jawaban"
                )

            if not has_correct:
                errors.append(
                    f"Soal No. {idx} (Pilihan Ganda): Kunci jawaban belum dipilih"
                )

        elif q_type == "true_false":
            if len(q.options) < 2:
                errors.append(
                    f"Soal No. {idx} (Benar/Salah): Opsi Benar dan Salah belum lengkap"
                )
            if not any(opt.is_correct for opt in q.options):
                errors.append(
                    f"Soal No. {idx} (Benar/Salah): Kunci jawaban belum dipilih"
                )

        elif q_type == "short_answer":
            require_manual = bool(settings.get("require_manual_grading", False))
            acceptable_answers = settings.get("acceptable_answers", []) or []
            first_key = (acceptable_answers[0] if acceptable_answers else "").strip()
            if not require_manual and not first_key:
                errors.append(
                    f"Soal No. {idx} (Isian Singkat): Kunci jawaban belum diisi"
                )

        elif q_type == "multiple_choice_complex":
            if not q.image_url and not (q.stimulus or "").strip():
                errors.append(
                    f"Soal No. {idx} (PG Kompleks): Stimulus/bacaan wajib diisi"
                )

            if pgk_type == "table_validation":
                statements = settings.get("statements", []) or []
                statement_answers = settings.get("statement_answers", []) or []
                valid_statements = [s for s in statements if (s or "").strip()]
                has_image_mode = bool(q.image_url)

                if not has_image_mode and len(valid_statements) < 2:
                    errors.append(
                        f"Soal No. {idx} (PG Kompleks): Minimal harus ada 2 pernyataan"
                    )
                elif has_image_mode and len(valid_statements) < 2 and len(statement_answers) < 2:
                    errors.append(
                        f"Soal No. {idx} (PG Kompleks): Minimal harus ada 2 pernyataan"
                    )

                required_answers_count = len(valid_statements)
                if has_image_mode:
                    required_answers_count = max(required_answers_count, 2)

                if len(statement_answers) < required_answers_count:
                    errors.append(
                        f"Soal No. {idx} (PG Kompleks): Jawaban Benar/Salah pernyataan belum lengkap"
                    )
            else:
                real_options_count = _count_real_options(q.options)
                has_embedded_options = _has_embedded_choice_lines(q.question_text)
                is_image_mode = bool(q.image_url)
                correct_count = sum(1 for opt in q.options if opt.is_correct)
                permissive_key_only_mode = correct_count >= 2 and not is_image_mode and not has_embedded_options

                if real_options_count < PGK_CHECKBOX_MIN_OPTIONS and not is_image_mode and not has_embedded_options and not permissive_key_only_mode:
                    errors.append(
                        f"Soal No. {idx} (PG Kompleks): Minimal harus ada {PGK_CHECKBOX_MIN_OPTIONS} opsi jawaban"
                    )

                if correct_count < 2:
                    errors.append(
                        f"Soal No. {idx} (PG Kompleks): Minimal 2 kunci jawaban harus dicentang"
                    )

    if errors:
        bullet_list = "\n".join(f"- {item}" for item in errors)
        raise HTTPException(
            status_code=400,
            detail=(
                "Ujian belum siap dipublish:\n"
                f"{bullet_list}\n"
                "Silakan lengkapi soal-soal tersebut terlebih dahulu."
            )
        )


# ============== PUBLIC ENDPOINTS ==============

@router.get("/results/all", response_model=List[ExamResponse])
async def get_exams_with_results(
    current_user: User = Depends(get_current_exam_monitor),
    db: AsyncSession = Depends(get_db_read)
):
    """
    Get all COMPLETED exams for results page (excluding deleted exams).

    Shows exams where:
    1. Completed + never had results (fresh exams for monitoring)
    2. OR has submitted/completed sessions (exams with actual results)

    IMPORTANT: Excludes soft-deleted exams AND exams that had results deleted

    Logic:
    - Fresh exam (has_ever_had_results=False) → SHOW (monitoring)
    - Exam with results (has sessions) → SHOW (view results)
    - Exam with deleted results (has_ever_had_results=True, no sessions) → HIDE
    - Deleted exam (is_deleted=True) → HIDE

    Authorization: Teacher/Admin only
    """
    now = datetime.now(timezone.utc)

    # Get exam IDs that currently have results
    exams_with_results_subquery = (
        select(ExamSession.exam_id)
        .where(ExamSession.status.in_(["submitted", "completed"]))
        .distinct()
    )

    # Main query: Get exams that are either:
    # 1. Completed + never had results (fresh exams)
    # 2. Currently have results
    # Exclude: soft-deleted exams + exams with deleted results
    query = (
        select(Exam)
        .options(
            noload("*"),
            joinedload(Exam.creator).noload("*"),
            noload(Exam.questions),
            noload(Exam.sessions),
            noload(Exam.schedules),
        )
        .where(
            or_(
                # Fresh completed exams (never had results) - for monitoring
                and_(
                    Exam.end_time < now,
                    Exam.has_ever_had_results == False
                ),
                # Exams with current results
                Exam.id.in_(exams_with_results_subquery)
            )
        )
        .where(Exam.is_deleted == False)  # Exclude soft-deleted exams
        .order_by(Exam.created_at.desc())
    )

    if current_user.role == "teacher":
        query = query.where(Exam.creator_id == current_user.id)
    elif is_pengawas_user(current_user):
        query = query.where(Exam.creator.has(User.role != ROLE_DEVELOPER))

    # FIX: Only show published exams
    query = query.where(Exam.is_published == True)

    result = await db.execute(query)
    exams = result.scalars().all()

    question_count_map: Dict[int, int] = {}
    exam_ids = [exam.id for exam in exams]
    if exam_ids:
        count_rows = await db.execute(
            select(Question.exam_id, func.count(Question.id))
            .where(Question.exam_id.in_(exam_ids))
            .group_by(Question.exam_id)
        )
        question_count_map = {exam_id: total for exam_id, total in count_rows.all()}

    for exam in exams:
        setattr(exam, "_question_count", int(question_count_map.get(exam.id, 0)))

    return [ExamResponse.from_orm_with_wib(exam) for exam in exams]


@router.get("/my-results", response_model=List[ExamResponse])
async def get_my_exam_results(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_read)
):
    """Get current user's exam history and results."""
    # Only participant roles can access this endpoint.
    if not _is_exam_participant_role(current_user.role):
        raise HTTPException(
            status_code=403,
            detail="Hanya peserta ujian yang dapat melihat riwayat ujian sendiri",
        )

    result = await db.execute(
        select(ExamSession)
        .options(selectinload(ExamSession.exam))
        .where(
            ExamSession.user_id == current_user.id,
            ExamSession.status.in_(['completed', 'submitted'])
        )
        .order_by(ExamSession.end_time.desc())
    )
    sessions = result.scalars().all()

    # Map sessions to ExamResponse format (simplified for list view)
    # We return the exam data but enriched with session score
    history = []
    for session in sessions:
        if not session.exam:
            continue

        # Use archived metadata if exam was deleted
        exam_data = session.exam

        # Note: We are returning ExamResponse, so we map session data to it where possible
        # or rely on the frontend to call separate details endpoint if needed.
        # But for "My Results", ideally we'd have a specific schema.
        # Re-using ExamResponse for now as requested by Audit Report structure implies list logic.

        # BUT, standard ExamResponse doesn't have "my_score".
        # However, the audit recommendation just said "return results".
        # Let's map it to the Exam schema structure.

        history.append(ExamResponse.from_orm_with_wib(exam_data))

    return history


@router.get("/{exam_id}/participation-summary")
async def get_exam_participation_summary(
    exam_id: int,
    current_user: User = Depends(get_current_exam_monitor),
    db: AsyncSession = Depends(get_db_read),
):
    """
    Summarize target participants vs actual sessions.

    This separates:
    - target students who never started
    - target students who started but have not submitted
    - submitted sessions outside current target metadata
    """
    exam_result = await db.execute(
        select(
            Exam.id,
            Exam.title,
            Exam.creator_id,
            Exam.allowed_classes,
            Exam.allowed_students,
            User.role.label("creator_role"),
        )
        .join(User, User.id == Exam.creator_id)
        .where(Exam.id == exam_id, Exam.is_deleted == False)
    )
    exam = exam_result.mappings().one_or_none()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")

    _enforce_developer_exam_visibility(current_user, exam["creator_role"])
    if (
        int(exam["creator_id"]) != current_user.id
        and not current_user.is_admin
        and not is_pengawas_user(current_user)
    ):
        raise HTTPException(status_code=403, detail="Not authorized to view this exam's participation summary")

    allowed_class_values = [
        value.strip()
        for value in (exam["allowed_classes"] or "").split(",")
        if value.strip()
    ]
    allowed_student_values = [
        value.strip()
        for value in (exam["allowed_students"] or "").split(",")
        if value.strip()
    ]

    target_conditions = [User.role == "student", User.is_active == True]
    restriction_clauses = []
    if allowed_class_values:
        restriction_clauses.append(User.student_class.in_(allowed_class_values))
    if allowed_student_values:
        restriction_clauses.append(User.id.cast(sqlalchemy.String).in_(allowed_student_values))
    if restriction_clauses:
        target_conditions.append(or_(*restriction_clauses))

    target_result = await db.execute(
        select(User.id, User.username, User.full_name, User.student_class)
        .where(*target_conditions)
        .order_by(User.student_class.asc(), User.full_name.asc(), User.username.asc())
    )
    target_rows = [dict(row) for row in target_result.mappings().all()]
    target_ids = {int(row["id"]) for row in target_rows}
    target_by_id = {int(row["id"]): row for row in target_rows}

    sessions_result = await db.execute(
        select(
            ExamSession.id.label("session_id"),
            ExamSession.user_id,
            ExamSession.status,
            ExamSession.start_time,
            ExamSession.end_time,
            User.username,
            User.full_name,
            User.student_class,
        )
        .join(User, User.id == ExamSession.user_id)
        .where(ExamSession.exam_id == exam_id)
        .order_by(ExamSession.start_time.desc(), ExamSession.id.desc())
    )
    session_rows = [dict(row) for row in sessions_result.mappings().all()]

    sessions_by_user: Dict[int, List[Dict[str, Any]]] = {}
    for row in session_rows:
        sessions_by_user.setdefault(int(row["user_id"]), []).append(row)

    submitted_statuses = {"submitted", "completed"}
    submitted_user_ids = {
        int(row["user_id"])
        for row in session_rows
        if str(row.get("status") or "").lower() in submitted_statuses
    }
    users_with_any_session = set(sessions_by_user.keys())

    target_submitted_ids = target_ids & submitted_user_ids
    not_started_ids = target_ids - users_with_any_session
    started_not_submitted_ids = {
        user_id
        for user_id in (target_ids & users_with_any_session)
        if user_id not in submitted_user_ids
    }
    submitted_outside_target_ids = submitted_user_ids - target_ids

    non_submitted_status_counts: Dict[str, int] = {}
    for user_id in started_not_submitted_ids:
        latest_status = str((sessions_by_user.get(user_id) or [{}])[0].get("status") or "unknown")
        non_submitted_status_counts[latest_status] = non_submitted_status_counts.get(latest_status, 0) + 1

    outside_students = []
    for user_id in sorted(submitted_outside_target_ids):
        latest = (sessions_by_user.get(user_id) or [{}])[0]
        outside_students.append(
            _student_summary_entry(
                {
                    "id": user_id,
                    "username": latest.get("username") or "",
                    "full_name": latest.get("full_name") or latest.get("username") or "",
                    "student_class": latest.get("student_class") or "",
                }
            )
        )

    return {
        "exam_id": exam_id,
        "exam_title": exam["title"],
        "restrictions": {
            "allowed_classes": allowed_class_values,
            "allowed_students_count": len(allowed_student_values),
        },
        "target_count": len(target_ids),
        "submitted_in_target_count": len(target_submitted_ids),
        "submitted_total_count": len(submitted_user_ids),
        "submitted_outside_target_count": len(submitted_outside_target_ids),
        "not_started_count": len(not_started_ids),
        "started_not_submitted_count": len(started_not_submitted_ids),
        "session_total_count": len(session_rows),
        "non_submitted_status_counts": non_submitted_status_counts,
        "not_started_students": [
            _student_summary_entry(target_by_id[user_id])
            for user_id in sorted(not_started_ids)
        ][:200],
        "started_not_submitted_students": [
            _student_summary_entry(target_by_id[user_id])
            for user_id in sorted(started_not_submitted_ids)
        ][:200],
        "submitted_outside_target_students": outside_students[:200],
    }


@router.get("/{exam_id}/participation-summary/export")
async def export_exam_participation_summary(
    exam_id: int,
    format: str = Query(
        "csv",
        pattern="^(csv|excel|xls|pdf|docx|word)$",
        description="Export format: csv, excel/xls, pdf, docx/word",
    ),
    current_user: User = Depends(get_current_exam_monitor),
    db: AsyncSession = Depends(get_db_read),
):
    """Export target-vs-submission participation rows as CSV, Excel, PDF, or Word."""
    require_feature_enabled(
        settings.heavy_exports_active,
        "heavy_export",
        status_code=503,
        message="Ekspor partisipasi sedang dinonaktifkan selama mode ujian/puncak.",
    )
    exam_result = await db.execute(
        select(
            Exam.id,
            Exam.title,
            Exam.subject,
            Exam.exam_type,
            Exam.creator_id,
            Exam.allowed_classes,
            Exam.allowed_students,
            User.role.label("creator_role"),
            User.full_name.label("creator_name"),
            User.username.label("creator_username"),
        )
        .join(User, User.id == Exam.creator_id)
        .where(Exam.id == exam_id, Exam.is_deleted == False)
    )
    exam = exam_result.mappings().one_or_none()
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")

    _enforce_developer_exam_visibility(current_user, exam["creator_role"])
    if (
        int(exam["creator_id"]) != current_user.id
        and not current_user.is_admin
        and not is_pengawas_user(current_user)
    ):
        raise HTTPException(status_code=403, detail="Not authorized to export this exam's participation summary")

    allowed_class_values = [
        value.strip()
        for value in (exam["allowed_classes"] or "").split(",")
        if value.strip()
    ]
    allowed_student_values = [
        value.strip()
        for value in (exam["allowed_students"] or "").split(",")
        if value.strip()
    ]

    target_conditions = [User.role == "student", User.is_active == True]
    restriction_clauses = []
    if allowed_class_values:
        restriction_clauses.append(User.student_class.in_(allowed_class_values))
    if allowed_student_values:
        restriction_clauses.append(User.id.cast(sqlalchemy.String).in_(allowed_student_values))
    if restriction_clauses:
        target_conditions.append(or_(*restriction_clauses))

    target_result = await db.execute(
        select(User.id, User.username, User.full_name, User.student_class)
        .where(*target_conditions)
        .order_by(User.student_class.asc(), User.full_name.asc(), User.username.asc())
    )
    target_rows = [dict(row) for row in target_result.mappings().all()]
    target_by_id = {int(row["id"]): row for row in target_rows}
    target_ids = set(target_by_id.keys())

    sessions_result = await db.execute(
        select(
            ExamSession.id.label("session_id"),
            ExamSession.user_id,
            ExamSession.status,
            ExamSession.start_time,
            ExamSession.end_time,
            ExamSession.score,
            User.username,
            User.full_name,
            User.student_class,
        )
        .join(User, User.id == ExamSession.user_id)
        .where(ExamSession.exam_id == exam_id)
        .order_by(ExamSession.start_time.desc(), ExamSession.id.desc())
    )
    session_rows = [dict(row) for row in sessions_result.mappings().all()]

    sessions_by_user: Dict[int, List[Dict[str, Any]]] = {}
    for row in session_rows:
        sessions_by_user.setdefault(int(row["user_id"]), []).append(row)

    submitted_statuses = {"submitted", "completed"}
    submitted_user_ids = {
        int(row["user_id"])
        for row in session_rows
        if str(row.get("status") or "").lower() in submitted_statuses
    }

    def latest_session(user_id: int, *, submitted_only: bool = False) -> Optional[Dict[str, Any]]:
        for session_row in sessions_by_user.get(user_id, []):
            if not submitted_only or str(session_row.get("status") or "").lower() in submitted_statuses:
                return session_row
        return None

    target_with_session_ids = target_ids & set(sessions_by_user.keys())
    summary_rows = [
        ("Exam ID", exam_id),
        ("Nama Ujian", exam["title"]),
        ("Mata Pelajaran", exam["subject"] or "-"),
        ("Jenis Ujian", exam["exam_type"] or "-"),
        ("Guru/Admin", exam["creator_name"] or exam["creator_username"] or "-"),
        ("Kelas Target", ", ".join(allowed_class_values) or "Semua/khusus peserta"),
        ("Jumlah Peserta Khusus", len(allowed_student_values)),
        ("Target", len(target_ids)),
        ("Submitted Total", len(submitted_user_ids)),
        ("Submitted Dalam Target", len(target_ids & submitted_user_ids)),
        ("Belum Start", len(target_ids - set(sessions_by_user.keys()))),
        ("Sudah Start Belum Submit", len(target_with_session_ids - submitted_user_ids)),
        ("Submitted Di Luar Target", len(submitted_user_ids - target_ids)),
    ]
    table_headers = [
        "No",
        "Kategori",
        "User ID",
        "Username",
        "Nama",
        "Kelas",
        "Session ID",
        "Status Session",
        "Nilai",
        "Mulai",
        "Selesai/Submit",
        "Keterangan",
    ]
    table_rows: List[List[Any]] = []

    row_no = 1
    sorted_target_ids = sorted(
        target_ids,
        key=lambda user_id: (
            str(target_by_id[user_id].get("student_class") or ""),
            str(target_by_id[user_id].get("full_name") or target_by_id[user_id].get("username") or ""),
        ),
    )
    for user_id in sorted_target_ids:
        target = target_by_id[user_id]
        session = latest_session(user_id, submitted_only=user_id in submitted_user_ids)
        if user_id in submitted_user_ids:
            category = "submitted"
            note = "Sudah submit"
        elif session:
            category = "started_not_submitted"
            note = "Sudah ada sesi tetapi belum submitted"
        else:
            category = "not_started"
            note = "Belum pernah membuat sesi ujian"
        table_rows.append([
            row_no,
            category,
            user_id,
            target.get("username") or "",
            target.get("full_name") or target.get("username") or "",
            target.get("student_class") or "",
            session.get("session_id") if session else "",
            session.get("status") if session else "",
            str(session.get("score")) if session and session.get("score") is not None else "",
            _format_wib_datetime(session.get("start_time")) if session else "",
            _format_wib_datetime(session.get("end_time")) if session else "",
            note,
        ])
        row_no += 1

    for user_id in sorted(submitted_user_ids - target_ids):
        session = latest_session(user_id, submitted_only=True) or {}
        table_rows.append([
            row_no,
            "submitted_outside_target",
            user_id,
            session.get("username") or "",
            session.get("full_name") or session.get("username") or "",
            session.get("student_class") or "",
            session.get("session_id") or "",
            session.get("status") or "",
            str(session.get("score")) if session.get("score") is not None else "",
            _format_wib_datetime(session.get("start_time")),
            _format_wib_datetime(session.get("end_time")),
            "Submitted, tetapi di luar target metadata saat ini",
        ])
        row_no += 1

    export_format = format.lower().strip()
    if export_format == "word":
        export_format = "docx"
    if export_format == "excel":
        export_format = "xls"

    safe_title = re.sub(r"[^\w\s-]", "", exam["title"] or "ujian").strip()
    safe_title = re.sub(r"\s+", "_", safe_title) or "ujian"
    filename_base = f"kehadiran_ujian_{safe_title}_{datetime.now().strftime('%Y%m%d')}"

    if export_format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["Laporan Kehadiran Ujian"])
        for label, value in summary_rows:
            writer.writerow([label, value])
        writer.writerow([])
        writer.writerow(table_headers)
        writer.writerows(table_rows)
        return Response(
            content="\ufeff" + output.getvalue(),
            media_type="text/csv; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.csv"'},
        )

    if export_format == "xls":
        def td(value: Any, *, header: bool = False) -> str:
            tag = "th" if header else "td"
            return f"<{tag}>{html.escape(str(value if value is not None else ''))}</{tag}>"

        summary_html = "".join(
            f"<tr>{td(label, header=True)}{td(value)}</tr>" for label, value in summary_rows
        )
        header_html = "".join(td(header, header=True) for header in table_headers)
        body_html = "".join(
            "<tr>" + "".join(td(value) for value in row) + "</tr>" for row in table_rows
        )
        html_content = f"""
        <html xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:x="urn:schemas-microsoft-com:office:excel" xmlns="http://www.w3.org/TR/REC-html40">
        <head><meta charset="utf-8"><style>
            body {{ font-family: Arial, sans-serif; }}
            h2 {{ color: #0f172a; }}
            table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
            th {{ background: #0f172a; color: white; font-weight: bold; }}
            th, td {{ border: 1px solid #000; padding: 6px; vertical-align: top; }}
        </style></head>
        <body>
            <h2>Laporan Kehadiran Ujian</h2>
            <table>{summary_html}</table>
            <table><thead><tr>{header_html}</tr></thead><tbody>{body_html}</tbody></table>
        </body></html>
        """
        return Response(
            content="\ufeff" + html_content,
            media_type="application/vnd.ms-excel; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.xls"'},
        )

    if export_format == "pdf":
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        except Exception as exc:
            raise HTTPException(status_code=501, detail="PDF export tidak tersedia. Install ReportLab.") from exc

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), rightMargin=18, leftMargin=18, topMargin=18, bottomMargin=18)
        styles = getSampleStyleSheet()
        story = [Paragraph("Laporan Kehadiran Ujian", styles["Title"]), Spacer(1, 8)]
        summary_table = Table([[str(label), str(value)] for label, value in summary_rows], colWidths=[150, 420])
        summary_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.extend([summary_table, Spacer(1, 10)])
        pdf_headers = ["No", "Kategori", "Username", "Nama", "Kelas", "Status", "Nilai", "Keterangan"]
        pdf_rows = [pdf_headers]
        for row in table_rows:
            pdf_rows.append([row[0], row[1], row[3], row[4], row[5], row[7], row[8], row[11]])
        pdf_table = Table(pdf_rows, repeatRows=1, colWidths=[28, 82, 82, 150, 48, 62, 38, 210])
        pdf_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 6),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ]))
        story.append(pdf_table)
        doc.build(story)
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.pdf"'},
        )

    if export_format == "docx":
        try:
            from docx import Document
            from docx.enum.section import WD_ORIENT
            from docx.shared import Inches
        except Exception as exc:
            raise HTTPException(status_code=501, detail="Word export tidak tersedia. Install python-docx.") from exc

        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        document.add_heading("Laporan Kehadiran Ujian", level=1)
        for label, value in summary_rows:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(str(value))
        document.add_paragraph("")
        table = document.add_table(rows=1, cols=len(table_headers))
        table.style = "Table Grid"
        for idx, header in enumerate(table_headers):
            table.rows[0].cells[idx].text = header
        for row in table_rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = str(value if value is not None else "")
        buffer = io.BytesIO()
        document.save(buffer)
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.docx"'},
        )

    raise HTTPException(status_code=400, detail="Format export tidak didukung")


@router.get("/{exam_id}/results")
async def get_exam_results(
    exam_id: int,
    include_breakdown: bool = Query(
        False,
        description="Include per-question score breakdown in response payload",
    ),
    current_user: User = Depends(get_current_exam_monitor),
    db: AsyncSession = Depends(get_db_read)
):
    """
    Get all results/sessions for a specific exam.

    Returns list of completed/submitted sessions with student info and scores.
    """
    if is_pengawas_user(current_user) and include_breakdown:
        raise HTTPException(
            status_code=403,
            detail="Pengawas tidak diizinkan melihat rincian soal atau kunci jawaban.",
        )
    # Verify exam exists and user has access
    exam_result = await db.execute(
        select(
            Exam.id,
            Exam.title,
            Exam.subject,
            Exam.exam_type,
            Exam.passing_score,
            Exam.creator_id,
            Exam.is_published,
            User.role.label("creator_role"),
        )
        .join(User, User.id == Exam.creator_id)
        .where(Exam.id == exam_id, Exam.is_deleted == False)
    )
    exam = exam_result.mappings().one_or_none()

    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")

    _enforce_developer_exam_visibility(current_user, exam["creator_role"])

    if is_pengawas_user(current_user):
        if not exam["is_published"]:
            raise HTTPException(status_code=404, detail="Exam not found")
    elif exam["creator_id"] != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=403, detail="Not authorized to view this exam's results")

    viewer_scope = _build_exam_results_viewer_scope(current_user, int(exam["creator_id"]))
    cache_key = _build_exam_results_cache_key(exam_id, include_breakdown, viewer_scope)
    cached_results = await _get_cached_exam_results(cache_key)
    if cached_results is not None:
        return cached_results

    # Get all completed/submitted sessions with student info
    try:
        sessions_result = await db.execute(
            select(
                ExamSession.id.label("session_id"),
                ExamSession.user_id,
                ExamSession.start_time,
                ExamSession.end_time,
                ExamSession.score,
                ExamSession.violation_count,
                ExamSession.status,
                User.id.label("student_id"),
                User.full_name,
                User.username,
                User.student_class,
            )
            .join(User, ExamSession.user_id == User.id)
            .where(
                ExamSession.exam_id == exam_id,
                ExamSession.status.in_(["submitted", "completed"])
            )
            .order_by(ExamSession.end_time.desc(), ExamSession.id.desc())
        )
        session_rows = [dict(row) for row in sessions_result.mappings().all()]
        if not session_rows:
            await _set_cached_exam_results(cache_key, [], include_breakdown=include_breakdown)
            return []
        selected_session_rows = _pick_latest_scored_exam_session_per_user(session_rows)

        ordered_questions: List[Dict[str, Any]] = []
        answers_by_session: Dict[int, Dict[int, Dict[str, Any]]] = {}

        if include_breakdown:
            questions_result = await db.execute(
                select(
                    Question.id.label("question_id"),
                    Question.question_type,
                    Question.points,
                    Question.order_index,
                )
                .where(Question.exam_id == exam_id)
                .order_by(Question.order_index.asc(), Question.id.asc())
            )
            ordered_questions = [dict(row) for row in questions_result.mappings().all()]

            if ordered_questions:
                session_ids = [int(row["session_id"]) for row in selected_session_rows]
                answers_result = await db.execute(
                    select(
                        Answer.session_id,
                        Answer.question_id,
                        Answer.is_correct,
                        Answer.points_earned,
                    )
                    .where(Answer.session_id.in_(session_ids))
                )
                for answer_row in answers_result.mappings().all():
                    session_answer_map = answers_by_session.setdefault(
                        int(answer_row["session_id"]),
                        {},
                    )
                    session_answer_map[int(answer_row["question_id"])] = {
                        "is_correct": answer_row["is_correct"],
                        "points_earned": answer_row["points_earned"],
                    }

        passing_score = float(exam["passing_score"]) if exam["passing_score"] is not None else 70.0
        results: List[Dict[str, Any]] = []

        for row in selected_session_rows:
            start_time = row["start_time"]
            end_time = row["end_time"]
            duration_seconds = 0
            if start_time and end_time:
                duration_seconds = int((end_time - start_time).total_seconds())

            score_value = float(row["score"]) if row["score"] is not None else 0.0
            passed = score_value >= passing_score if row["score"] is not None else False

            session_metadata: Dict[str, Any] = {}
            if include_breakdown and ordered_questions:
                score_breakdown = []
                answers_by_question_id = answers_by_session.get(int(row["session_id"]), {})

                for question in ordered_questions:
                    answer = answers_by_question_id.get(int(question["question_id"]))
                    max_points = float(question["points"]) if question["points"] is not None else 0.0
                    points_earned = None
                    if answer and answer["points_earned"] is not None:
                        points_earned = float(answer["points_earned"])

                    if answer is None:
                        item_status = "not_answered"
                        item_is_correct = False
                        points_earned = 0.0
                    elif points_earned is None:
                        item_status = "pending"
                        item_is_correct = None
                    elif question["question_type"] in ("essay", "short_answer"):
                        if points_earned >= max_points:
                            item_status = "correct"
                            item_is_correct = True
                        elif points_earned <= 0:
                            item_status = "incorrect"
                            item_is_correct = False
                        else:
                            item_status = "partial"
                            item_is_correct = None
                    else:
                        if answer["is_correct"] is True:
                            item_status = "correct"
                            item_is_correct = True
                        elif answer["is_correct"] is False and points_earned > 0:
                            item_status = "partial"
                            item_is_correct = False
                        elif answer["is_correct"] is False:
                            item_status = "incorrect"
                            item_is_correct = False
                        elif points_earned > 0:
                            item_status = "partial"
                            item_is_correct = None
                        else:
                            item_status = "incorrect"
                            item_is_correct = False

                    score_breakdown.append(
                        {
                            "question_id": str(question["question_id"]),
                            "question_type": question["question_type"],
                            "points_earned": points_earned,
                            "max_points": max_points,
                            "is_correct": item_is_correct,
                            "status": item_status,
                        }
                    )

                session_metadata["score_breakdown"] = score_breakdown

            results.append(
                {
                    "id": int(row["session_id"]),
                    "user_id": int(row["user_id"]),
                    "user": {
                        "id": int(row["student_id"]),
                        "full_name": row["full_name"] or row["username"] or "Unknown",
                        "username": row["username"] or "unknown",
                        "student_class": row["student_class"] or "",
                    },
                    "exam": {
                        "id": int(exam["id"]),
                        "title": exam["title"],
                        "subject": exam["subject"] or "",
                        "exam_type": exam["exam_type"] or "",
                        "passing_score": passing_score,
                    },
                    "score": score_value,
                    "start_time": start_time,
                    "end_time": end_time,
                    "submitted_at": end_time,
                    "duration_seconds": duration_seconds,
                    "violation_count": row["violation_count"] or 0,
                    "passed": passed,
                    "status": row["status"] or "unknown",
                    "session_metadata": session_metadata,
                }
            )

        await _set_cached_exam_results(cache_key, results, include_breakdown=include_breakdown)
        return results
    except Exception:
        logger.exception("Error in get_exam_results")
        raise HTTPException(status_code=500, detail="Terjadi kesalahan pada server saat memuat hasil ujian.")


@router.get("/{exam_id}/sessions/{session_id}/review")
async def get_session_answer_review(
    exam_id: int,
    session_id: int,
    current_user: User = Depends(get_current_teacher),
    db: AsyncSession = Depends(get_db_read),
):
    """
    Fetch detailed answer review payload for admin/teacher result inspection.
    Lightweight by using 3 bounded queries:
    - session + participant + exam metadata
    - ordered questions + options
    - submitted answers for this session
    """
    if is_pengawas_user(current_user):
        raise HTTPException(
            status_code=403,
            detail="Pengawas tidak diizinkan meninjau jawaban.",
        )
    session_row_result = await db.execute(
        select(
            ExamSession.id.label("session_id"),
            ExamSession.user_id,
            ExamSession.score,
            ExamSession.status,
            ExamSession.start_time,
            ExamSession.end_time,
            User.id.label("student_id"),
            User.full_name,
            User.username,
            User.student_class,
            Exam.id.label("exam_id"),
            Exam.title,
            Exam.subject,
            Exam.exam_type,
            Exam.passing_score,
            Exam.creator_id,
        )
        .join(User, ExamSession.user_id == User.id)
        .join(Exam, ExamSession.exam_id == Exam.id)
        .where(
            ExamSession.id == session_id,
            ExamSession.exam_id == exam_id,
            Exam.is_deleted == False,
        )
    )
    session_row = session_row_result.mappings().one_or_none()
    if not session_row:
        raise HTTPException(status_code=404, detail="Sesi ujian tidak ditemukan")

    creator_role = await _get_exam_creator_role(db, int(session_row["creator_id"]))
    _enforce_developer_exam_visibility(current_user, creator_role)

    if (
        int(session_row["creator_id"]) != current_user.id
        and not bool(current_user.is_admin)
    ):
        raise HTTPException(status_code=403, detail="Tidak memiliki akses")

    questions_result = await db.execute(
        select(Question)
        .options(selectinload(Question.options))
        .where(Question.exam_id == exam_id)
        .order_by(Question.order_index.asc(), Question.id.asc())
    )
    questions = questions_result.scalars().all()

    answers_result = await db.execute(
        select(Answer)
        .where(Answer.session_id == session_id)
    )
    answers = answers_result.scalars().all()
    answers_by_question = {int(a.question_id): a for a in answers}

    total_questions = len(questions)
    correct_count = 0
    partial_count = 0
    incorrect_count = 0
    pending_count = 0
    unanswered_count = 0
    review_items: List[Dict[str, Any]] = []

    for order_idx, question in enumerate(questions, start=1):
        max_points = float(question.points or 0.0)
        answer = answers_by_question.get(int(question.id))
        question_settings = dict(question.question_settings or {})
        option_map = _build_option_map(question.options or [])
        options_payload = list(option_map.values())
        correct_options = [opt for opt in options_payload if opt["is_correct"]]

        status = _status_from_answer(answer, max_points)
        if status == "correct":
            correct_count += 1
        elif status == "partial":
            partial_count += 1
        elif status == "pending":
            pending_count += 1
        elif status == "not_answered":
            unanswered_count += 1
        else:
            incorrect_count += 1

        points_earned = float(answer.points_earned) if answer and answer.points_earned is not None else 0.0
        answer_meta = dict(answer.answer_metadata or {}) if answer else {}
        statement_answers = (
            answer_meta.get("statement_answers")
            if isinstance(answer_meta.get("statement_answers"), dict)
            else {}
        )

        student_answer_display = "-"
        answer_key_display = "-"
        student_answer_payload: Dict[str, Any] = {"type": question.question_type, "display": "-"}
        answer_key_payload: Dict[str, Any] = {"type": question.question_type, "display": "-"}

        if question.question_type in {"multiple_choice", "true_false"}:
            selected = option_map.get(int(answer.selected_option_id)) if answer and answer.selected_option_id else None
            key_opt = correct_options[0] if correct_options else None

            if selected:
                student_answer_display = f"{selected['label']}. {selected['text']}"
                student_answer_payload.update({
                    "selected_option_id": selected["id"],
                    "selected_option_label": selected["label"],
                    "selected_option_text": selected["text"],
                    "display": student_answer_display,
                })
            else:
                student_answer_payload["display"] = "-"

            if key_opt:
                answer_key_display = f"{key_opt['label']}. {key_opt['text']}"
                answer_key_payload.update({
                    "correct_option_id": key_opt["id"],
                    "correct_option_label": key_opt["label"],
                    "correct_option_text": key_opt["text"],
                    "display": answer_key_display,
                })
            else:
                answer_key_payload["display"] = "-"

        elif question.question_type == "multiple_choice_complex":
            pgk_type = question.pgk_type or str(question_settings.get("pgk_type") or "checkbox")
            if pgk_type == "table_validation":
                statements = _resolve_question_statements(question_settings)
                key_map = _resolve_statement_keys(question_settings, len(statements))
                table_rows: List[Dict[str, Any]] = []
                student_lines: List[str] = []
                key_lines: List[str] = []

                max_index = max(
                    len(statements),
                    len(key_map),
                    len(statement_answers) if isinstance(statement_answers, dict) else 0,
                )

                for idx in range(max_index):
                    key = str(idx)
                    statement_text = statements[idx] if idx < len(statements) else f"Pernyataan {idx + 1}"
                    correct_bool = key_map.get(key)
                    student_bool = _coerce_bool(statement_answers.get(key)) if isinstance(statement_answers, dict) else None
                    table_rows.append(
                        {
                            "index": idx + 1,
                            "statement": statement_text,
                            "student_answer": student_bool,
                            "correct_answer": correct_bool,
                        }
                    )
                    student_lines.append(
                        f"{idx + 1}. {statement_text}: "
                        f"{'Benar' if student_bool is True else ('Salah' if student_bool is False else '-')}"
                    )
                    key_lines.append(
                        f"{idx + 1}. {statement_text}: "
                        f"{'Benar' if correct_bool is True else ('Salah' if correct_bool is False else '-')}"
                    )

                student_answer_display = " | ".join(student_lines) if student_lines else "-"
                answer_key_display = " | ".join(key_lines) if key_lines else "-"
                student_answer_payload.update(
                    {
                        "pgk_type": "table_validation",
                        "rows": table_rows,
                        "display": student_answer_display,
                    }
                )
                answer_key_payload.update(
                    {
                        "pgk_type": "table_validation",
                        "rows": table_rows,
                        "display": answer_key_display,
                    }
                )
            else:
                selected_ids = [int(x) for x in (answer.selected_option_ids or [])] if answer else []
                selected_opts = [option_map[opt_id] for opt_id in selected_ids if opt_id in option_map]
                selected_opts.sort(key=lambda x: x["label"])

                student_answer_display = ", ".join(
                    f"{opt['label']}. {opt['text']}" for opt in selected_opts
                ) if selected_opts else "-"
                answer_key_display = ", ".join(
                    f"{opt['label']}. {opt['text']}" for opt in correct_options
                ) if correct_options else "-"

                student_answer_payload.update(
                    {
                        "pgk_type": "checkbox",
                        "selected_option_ids": [opt["id"] for opt in selected_opts],
                        "selected_options": selected_opts,
                        "display": student_answer_display,
                    }
                )
                answer_key_payload.update(
                    {
                        "pgk_type": "checkbox",
                        "correct_option_ids": [opt["id"] for opt in correct_options],
                        "correct_options": correct_options,
                        "display": answer_key_display,
                    }
                )

        elif question.question_type == "short_answer":
            acceptable_answers = question_settings.get("acceptable_answers") or []
            accepted_values = [
                str(item).strip()
                for item in acceptable_answers
                if str(item or "").strip()
            ]
            student_answer_display = str(answer.answer_text or "").strip() if answer else "-"
            answer_key_display = " / ".join(accepted_values) if accepted_values else "-"
            student_answer_payload.update(
                {
                    "answer_text": student_answer_display if student_answer_display != "-" else "",
                    "display": student_answer_display,
                }
            )
            answer_key_payload.update(
                {
                    "acceptable_answers": accepted_values,
                    "display": answer_key_display,
                }
            )

        else:  # essay + fallback text
            key_essay = (
                str(question_settings.get("answer_key") or "").strip()
                or str(question_settings.get("sample_answer") or "").strip()
            )
            student_answer_display = str(answer.answer_text or "").strip() if answer else "-"
            answer_key_display = key_essay or "-"
            student_answer_payload.update(
                {
                    "answer_text": student_answer_display if student_answer_display != "-" else "",
                    "display": student_answer_display,
                }
            )
            answer_key_payload.update(
                {
                    "sample_answer": key_essay,
                    "display": answer_key_display,
                }
            )

        review_items.append(
            {
                "question_id": int(question.id),
                "order_index": int(question.order_index if question.order_index is not None else order_idx),
                "question_number": order_idx,
                "question_type": question.question_type,
                "question_type_label": QUESTION_TYPE_LABELS.get(question.question_type, question.question_type),
                "question_text": question.question_text or "",
                "stimulus": question.stimulus or "",
                "max_points": max_points,
                "points_earned": points_earned,
                "status": status,
                "is_correct": answer.is_correct if answer else None,
                "answered_at": answer.answered_at.isoformat() if answer and answer.answered_at else None,
                "student_answer": student_answer_payload,
                "answer_key": answer_key_payload,
                "student_answer_display": student_answer_display,
                "answer_key_display": answer_key_display,
                "options": options_payload,
            }
        )

    review_items.sort(key=lambda item: (item["order_index"], item["question_id"]))
    score_value = float(session_row["score"] or 0.0)
    passing_score = float(session_row["passing_score"] or 70.0)

    return {
        "session_id": int(session_row["session_id"]),
        "exam": {
            "id": int(session_row["exam_id"]),
            "title": session_row["title"] or f"Ujian #{exam_id}",
            "subject": session_row["subject"] or "",
            "exam_type": session_row["exam_type"] or "",
            "passing_score": passing_score,
        },
        "student": {
            "id": int(session_row["student_id"]),
            "full_name": session_row["full_name"] or session_row["username"] or "Unknown",
            "username": session_row["username"] or "unknown",
            "student_class": session_row["student_class"] or "",
        },
        "session": {
            "status": session_row["status"] or "unknown",
            "start_time": session_row["start_time"].isoformat() if session_row["start_time"] else None,
            "end_time": session_row["end_time"].isoformat() if session_row["end_time"] else None,
            "score": score_value,
            "passed": score_value >= passing_score,
        },
        "summary": {
            "total_questions": total_questions,
            "answered_questions": total_questions - unanswered_count,
            "correct": correct_count,
            "partial": partial_count,
            "incorrect": incorrect_count,
            "pending": pending_count,
            "unanswered": unanswered_count,
        },
        "questions": review_items,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


@router.delete("/{exam_id}/results")
async def delete_exam_results(
    exam_id: int,
    current_user: User = Depends(get_current_teacher),
    db: AsyncSession = Depends(get_db)
):
    """
    Delete all results/sessions for an exam.

    WARNING: This permanently deletes student session data!
    Use with caution. Only accessible by exam creator or admin.
    """
    # Verify exam exists and user has access
    result = await db.execute(select(Exam).where(Exam.id == exam_id))
    exam = result.scalar_one_or_none()

    if not exam:
        raise HTTPException(status_code=404, detail="Ujian tidak ditemukan")

    await _enforce_exam_owner_or_admin_access(
        db,
        current_user,
        exam.creator_id,
    )

    # Delete all sessions and their answers for this exam
    delete_stmt = delete(ExamSession).where(ExamSession.exam_id == exam_id)
    delete_result = await db.execute(delete_stmt)
    await db.commit()
    await _invalidate_exam_results_cache(exam_id)

    deleted_count = delete_result.rowcount

    return {
        "success": True,
        "message": f"Berhasil menghapus {deleted_count} hasil ujian",
        "deleted_count": deleted_count
    }


# ============== TOKEN-BASED ACCESS ==============

class JoinExamRequest(BaseModel):
    token: str

class JoinExamResponse(BaseModel):
    exam_id: int
    title: str
    description: Optional[str]
    duration_minutes: int
    question_count: int
    allowed: bool
    message: str


@router.post("/join", response_model=JoinExamResponse)
async def join_exam_by_token(
    request: JoinExamRequest,
    raw_request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_read)
):
    """
    Join an exam using access token.

    Validates:
    1. Rate limit (5 attempts/min) - Anti Brute-force
    2. Token exists and is valid
    3. Exam is published
    4. Exam is within active time window
    5. Participant belongs to allowed class/list policy
    6. Student hasn't exceeded max attempts
    """
    # Rate limit check (5 per minute) with proxy-aware IP resolver.
    client_ip = get_client_ip(raw_request)
    is_allowed, remaining = await check_rate_limit(RateLimiters.JOIN_EXAM, f"{current_user.id}:{client_ip}")

    if not is_allowed:
        raise HTTPException(
            status_code=429,
            detail="Terlalu banyak percobaan token salah. Tunggu 1 menit.",
            headers={"Retry-After": "60"}
        )

    # Only exam participants can join exams.
    if not _is_exam_participant_role(current_user.role):
        raise HTTPException(
            status_code=403,
            detail="Hanya peserta ujian yang dapat mengikuti ujian",
        )

    # Normalize token (uppercase, trim)
    token = request.token.strip().upper()

    if len(token) != 6:
        raise HTTPException(status_code=400, detail="Token harus 6 karakter")

    # Find exam by token (keep query lightweight under burst join traffic).
    result = await db.execute(
        select(Exam, User.role.label("creator_role"))
        .options(noload("*"))
        .join(User, User.id == Exam.creator_id)
        .where(Exam.access_token == token)
    )
    exam_row = result.first()
    exam = exam_row[0] if exam_row else None
    exam_creator_role = exam_row[1] if exam_row else None

    if not exam:
        raise HTTPException(status_code=404, detail="Token ujian tidak valid")

    # Check if published
    if not exam.is_published:
        raise HTTPException(status_code=403, detail="Ujian belum dipublikasikan")

    # Check time window
    now = datetime.now(timezone.utc)
    if now < exam.start_time:
        raise HTTPException(status_code=403, detail="Ujian belum dimulai")
    if now > exam.end_time:
        raise HTTPException(status_code=403, detail="Ujian sudah berakhir")

    # Check participant restriction.
    _ensure_exam_participant_access(
        exam,
        current_user,
        exam_creator_role=exam_creator_role,
    )

    # Check max attempts (COUNT query avoids loading session rows)
    completed_attempts_result = await db.execute(
        select(func.count(ExamSession.id))
        .where(ExamSession.user_id == current_user.id)
        .where(ExamSession.exam_id == exam.id)
        .where(ExamSession.status.in_(("completed", "submitted")))
    )
    completed_attempts = int(completed_attempts_result.scalar() or 0)

    if completed_attempts >= exam.max_attempts:
        raise HTTPException(status_code=403, detail=f"Anda sudah menggunakan semua kesempatan ({exam.max_attempts}x)")

    question_count_result = await db.execute(
        select(func.count(Question.id)).where(Question.exam_id == exam.id)
    )
    question_count = int(question_count_result.scalar() or 0)

    return JoinExamResponse(
        exam_id=exam.id,
        title=exam.title,
        description=exam.description,
        duration_minutes=exam.duration_minutes,
        question_count=question_count,
        allowed=True,
        message="Token valid. Anda dapat memulai ujian."
    )


def _build_start_question_responses(
    questions_payload: List[Dict[str, Any]],
    *,
    exam_id: int,
    user_id: int,
    shuffle_questions: bool,
    shuffle_options: bool,
    secret_key: str,
) -> List[QuestionResponse]:
    questions_list: List[SimpleNamespace] = []
    for raw_question in questions_payload:
        raw_options = raw_question.get("options") or []
        normalized_options = [
            SimpleNamespace(
                id=safe_int(raw_option.get("id")) or 0,
                option_text=raw_option.get("option_text"),
                order_index=safe_int(raw_option.get("order_index")) or 0,
                option_group=raw_option.get("option_group") or "standard",
                pair_id=raw_option.get("pair_id"),
            )
            for raw_option in raw_options
        ]

        questions_list.append(
            SimpleNamespace(
                id=safe_int(raw_question.get("id")) or 0,
                question_text=raw_question.get("question_text"),
                stimulus=raw_question.get("stimulus"),
                question_type=raw_question.get("question_type"),
                pgk_type=raw_question.get("pgk_type"),
                difficulty_level=raw_question.get("difficulty_level"),
                question_settings=raw_question.get("question_settings") or {},
                points=float(raw_question.get("points") or 0),
                order_index=safe_int(raw_question.get("order_index")) or 0,
                image_url=raw_question.get("image_url"),
                video_url=raw_question.get("video_url"),
                audio_url=raw_question.get("audio_url"),
                cached_options=normalized_options,
            )
        )

    questions_list.sort(key=lambda q: q.order_index)

    if shuffle_questions:
        def get_question_hash(q_id: int) -> int:
            seed_str = f"{secret_key}_{user_id}_{exam_id}_question_{q_id}"
            return int(hashlib.md5(seed_str.encode()).hexdigest(), 16)

        questions_list.sort(key=lambda q: get_question_hash(q.id))

    questions: List[QuestionResponse] = []
    skipped_questions: List[Dict[str, Any]] = []

    for q in questions_list:
        try:
            question_settings = dict(q.question_settings or {})
            question_text = (q.question_text or "").strip()
            placeholder_source = str(
                question_settings.get("placeholder_source") or ""
            ).strip().lower()
            is_placeholder_question = _is_placeholder_question(question_settings)
            is_image_placeholder = (
                is_placeholder_question
                and bool(q.image_url)
                and placeholder_source == "image"
            )

            if not question_text:
                if is_image_placeholder:
                    question_text = (
                        "Perhatikan gambar soal berikut, lalu pilih jawaban yang benar."
                    )
                    logger.warning(
                        "EXAM_START | Question %s uses image-placeholder fallback text",
                        q.id,
                    )
                else:
                    logger.error("EXAM_START | Question %s has NO QUESTION_TEXT", q.id)
                    skipped_questions.append({"id": q.id, "reason": "no_text"})
                    continue

            q_settings = q.question_settings or {}
            pgk_type = q.pgk_type or q_settings.get("pgk_type", "checkbox")
            is_table_validation = (
                q.question_type == "multiple_choice_complex"
                and pgk_type == "table_validation"
            )

            requires_options = (
                q.question_type in ['multiple_choice', 'multiple_choice_complex', 'true_false']
                and not is_table_validation
            )
            question_options = list(getattr(q, "cached_options", []) or [])

            if requires_options and not question_options:
                logger.error(
                    "EXAM_START | Question %s (type: %s) has NO OPTIONS",
                    q.id,
                    q.question_type,
                )
                skipped_questions.append({
                    "id": q.id,
                    "text": (q.question_text or "")[:50],
                    "reason": "no_options",
                    "type": q.question_type,
                })
                continue

            options = []
            should_shuffle = bool(shuffle_options)

            if requires_options:
                options_list = sorted(question_options, key=lambda x: x.order_index)
                is_placeholder = is_placeholder_question
                can_shuffle_placeholder = _can_shuffle_placeholder_options(
                    question_settings,
                    has_image=bool(q.image_url)
                )

                if should_shuffle and (not is_placeholder or can_shuffle_placeholder):
                    seed_str = (
                        f"{secret_key}_{user_id}_{exam_id}_question_{q.id}_options"
                    )
                    options_list = _stable_shuffle_with_seed(options_list, seed_str)

                options = [
                    QuestionOptionResponse(
                        id=opt.id,
                        option_text=opt.option_text,
                        order_index=opt.order_index,
                        option_group=opt.option_group or "standard",
                        pair_id=opt.pair_id
                    )
                    for opt in options_list
                ]

            pgk_type = q.pgk_type or question_settings.get("pgk_type", "checkbox")
            is_table_validation = (
                q.question_type == "multiple_choice_complex"
                and pgk_type == "table_validation"
            )
            table_statement_shuffle_allowed = bool(
                question_settings.get("allow_table_statement_shuffle", True)
            ) if is_table_validation else False
            if is_table_validation:
                question_settings["allow_table_statement_shuffle"] = (
                    table_statement_shuffle_allowed
                )

            if should_shuffle and is_table_validation and table_statement_shuffle_allowed:
                statements = question_settings.get("statements", [])
                if statements:
                    normalized_texts = []
                    for s in statements:
                        if isinstance(s, dict):
                            text = str(s.get("text", "")).strip()
                        else:
                            text = str(s).strip()
                        normalized_texts.append(text)

                    informative_texts = [
                        t for t in normalized_texts
                        if t and t not in {"-", "--", "—", "–"}
                    ]
                    has_meaningful_statement_text = len(set(informative_texts)) >= 2
                    is_image_mode = bool(q.image_url)

                    if has_meaningful_statement_text and not is_image_mode:
                        indexed_stmts = [
                            {"text": s, "original_index": i}
                            for i, s in enumerate(statements)
                        ]
                        seed_str = (
                            f"{secret_key}_{user_id}_{exam_id}_question_{q.id}_statements"
                        )
                        indexed_stmts = _stable_shuffle_with_seed(indexed_stmts, seed_str)
                        question_settings["statements"] = indexed_stmts

            questions.append(QuestionResponse(
                id=q.id,
                question_text=question_text,
                stimulus=q.stimulus,
                question_type=q.question_type,
                pgk_type=q.pgk_type,
                difficulty_level=q.difficulty_level or "medium",
                category=None,
                tags=[],
                question_settings=question_settings,
                points=q.points,
                order_index=q.order_index,
                image_url=q.image_url,
                video_url=q.video_url,
                audio_url=q.audio_url,
                options=options
            ))

        except Exception as e:
            logger.error(
                "EXAM_START | Question %s FAILED to build: %s",
                q.id,
                str(e),
                exc_info=True,
            )
            skipped_questions.append({
                "id": q.id,
                "text": getattr(q, 'question_text', 'N/A')[:50],
                "reason": "exception",
                "error": str(e),
            })
            continue

    expected_count = len(questions_list)
    actual_count = len(questions)

    if skipped_questions:
        logger.error(
            "EXAM_START | SKIPPED %s questions: %s",
            len(skipped_questions),
            skipped_questions,
        )

    if actual_count < expected_count:
        logger.error(
            "EXAM_START | QUESTION COUNT MISMATCH! Expected %s, got %s",
            expected_count,
            actual_count,
        )
        raise HTTPException(
            status_code=500,
            detail=(
                f"Gagal memuat {expected_count - actual_count} soal dari ujian. "
                "Data ujian tidak lengkap. Silakan hubungi pengawas atau administrator."
            )
        )

    return questions


def _questions_to_start_payload(questions: Any) -> List[Dict[str, Any]]:
    payload: List[Dict[str, Any]] = []
    for question in questions:
        raw_options = list(getattr(question, "options", None) or [])
        payload.append(
            {
                "id": getattr(question, "id", 0),
                "question_text": getattr(question, "question_text", None),
                "stimulus": getattr(question, "stimulus", None),
                "question_type": getattr(question, "question_type", None),
                "pgk_type": getattr(question, "pgk_type", None),
                "difficulty_level": getattr(question, "difficulty_level", None),
                "question_settings": dict(getattr(question, "question_settings", None) or {}),
                "points": getattr(question, "points", 0),
                "order_index": getattr(question, "order_index", 0),
                "image_url": getattr(question, "image_url", None),
                "video_url": getattr(question, "video_url", None),
                "audio_url": getattr(question, "audio_url", None),
                "options": [
                    {
                        "id": getattr(option, "id", 0),
                        "option_text": getattr(option, "option_text", None),
                        "order_index": getattr(option, "order_index", 0),
                        "option_group": getattr(option, "option_group", None),
                        "pair_id": getattr(option, "pair_id", None),
                    }
                    for option in raw_options
                ],
            }
        )
    return payload


@router.post("/{exam_id}/start", response_model=ExamStartResponse)
async def start_exam_session(
    exam_id: int,
    request: Request,
    current_user: AuthenticatedUser = Depends(get_current_user_hot_path),
    db: AsyncSession = Depends(get_db)
):
    """Start an exam session (participant roles only)."""
    # Only participant roles can take exams.
    if not _is_exam_participant_role(current_user.role):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Hanya peserta ujian yang dapat mengikuti ujian"
        )

    async with bind_start_admission(request):
        await validate_seb_headers(request, exam_id, db, require_seb=True)
        exam_service = ExamService(db)
        async with start_db_segment("main"):
            exam = await exam_service.get_exam_start_projection(exam_id)

            if not exam:
                raise HTTPException(status_code=404, detail="Ujian tidak ditemukan")

            await _ensure_exam_start_option_integrity(db, exam_id)

            if not exam.is_published:
                raise HTTPException(status_code=400, detail="Ujian belum dipublikasikan")

            now = datetime.now(timezone.utc)
            if now < exam.start_time:
                raise HTTPException(status_code=400, detail="Ujian belum dimulai")
            if now > exam.end_time:
                raise HTTPException(status_code=400, detail="Ujian sudah berakhir")

            exam_creator_role = exam.creator.role if exam.creator else None
            _ensure_exam_participant_access(
                exam,
                current_user,
                exam_creator_role=exam_creator_role,
            )

            session_state = await exam_service.get_exam_start_session_state(
                current_user.id,
                exam_id,
            )
            completed_attempts = session_state.attempt_count
            if completed_attempts >= exam.max_attempts:
                raise HTTPException(status_code=400, detail="Batas percobaan sudah tercapai")

            existing_sessions = session_state.existing_sessions

            # Preload answer counts only for candidate resume sessions.
            answer_counts: Dict[int, int] = {}
            if len(existing_sessions) > 1:
                existing_session_ids = [s.id for s in existing_sessions]
                answer_count_result = await db.execute(
                    select(Answer.session_id, func.count(Answer.id))
                    .where(Answer.session_id.in_(existing_session_ids))
                    .group_by(Answer.session_id)
                )
                answer_counts = {int(sid): int(cnt or 0) for sid, cnt in answer_count_result.all()}

            # Check for resumable session.
            # If duplicate active sessions exist due reconnect/race, prefer the one with most saved answers.
            resumable_sessions = [s for s in existing_sessions if s.status in ("in_progress", "active")]
            is_resumed_session = False
            session = None
            if resumable_sessions:
                is_resumed_session = True
                resumable_sessions.sort(
                    key=lambda s: (
                        answer_counts.get(s.id, 0),
                        s.start_time or datetime.min.replace(tzinfo=timezone.utc),
                        s.id
                    ),
                    reverse=True
                )
                session = resumable_sessions[0]
                logger.info(
                    "EXAM_START | RESUME_SESSION | user=%s exam=%s session=%s answers=%s status=%s",
                    current_user.id,
                    exam_id,
                    session.id,
                    answer_counts.get(session.id, 0),
                    session.status
                )
            else:
                # Auto-reset terminated sessions only when cause is network/disconnection.
                recoverable_sessions = [
                    s for s in existing_sessions if s.status in ("terminated", "kicked")
                ]
                recoverable_sessions.sort(
                    key=lambda s: (
                        answer_counts.get(s.id, 0),
                        s.start_time or datetime.min.replace(tzinfo=timezone.utc),
                        s.id,
                    ),
                    reverse=True,
                )

                candidate_recoveries = []
                for candidate in recoverable_sessions:
                    logs_result = await db.execute(
                        select(ExamLog)
                        .options(noload("*"))
                        .where(ExamLog.session_id == candidate.id)
                        .order_by(ExamLog.created_at.desc(), ExamLog.id.desc())
                        .limit(30)
                    )
                    recovery = evaluate_session_recovery(candidate, logs_result.scalars().all())
                    if recovery.get("category") == RECOVERY_CATEGORY_ADMIN:
                        raise HTTPException(
                            status_code=409,
                            detail=(
                                "Sesi dihentikan oleh pengawas/admin. "
                                "Hubungi pengawas untuk membuka kembali sesi."
                            ),
                        )
                    candidate_recoveries.append((candidate, recovery))

                for candidate, recovery in candidate_recoveries:
                    if not recovery.get("allow_continue"):
                        continue

                    await db.execute(
                        update(ExamSession)
                        .where(ExamSession.id == candidate.id)
                        .values(
                            status="in_progress",
                            end_time=None,
                            terminated_by_admin=False,
                            emergency_exit_allowed=False,
                        )
                    )
                    candidate.status = "in_progress"
                    candidate.end_time = None
                    candidate.terminated_by_admin = False
                    candidate.emergency_exit_allowed = False
                    db.add(
                        ExamLog(
                            session_id=candidate.id,
                            event_type="SESSION_AUTO_RESET_NETWORK",
                            event_data={
                                "category": recovery.get("category"),
                                "message": recovery.get("message"),
                                "trigger": "start_exam_session",
                            },
                        )
                    )
                    session = candidate
                    is_resumed_session = True
                    logger.warning(
                        "EXAM_START | AUTO_RESET_SESSION | user=%s exam=%s session=%s category=%s",
                        current_user.id,
                        exam_id,
                        candidate.id,
                        recovery.get("category"),
                    )
                    break

            if session is None:
                # Create new session
                client_info = get_client_info(request)
                session = ExamSession(
                    user_id=current_user.id,
                    exam_id=exam_id,
                    start_time=now,
                    status="in_progress",
                    ip_address=client_info["ip_address"],
                    user_agent=client_info["user_agent"],
                    seb_detected=client_info["seb_detected"]
                )
                db.add(session)
                try:
                    await db.flush()
                    db.add(
                        ExamLog(
                            session_id=session.id,
                            event_type="SESSION_START",
                            event_data={
                                "ip": client_info["ip_address"],
                                "seb_detected": client_info["seb_detected"],
                                "exam_snapshot": {
                                    "title": exam.title,
                                    "subject": exam.subject,
                                    "exam_type": exam.exam_type,
                                    "allowed_classes": exam.allowed_classes,
                                    "allowed_students": exam.allowed_students,
                                    "start_time": exam.start_time.isoformat() if exam.start_time else None,
                                    "end_time": exam.end_time.isoformat() if exam.end_time else None,
                                    "duration_minutes": exam.duration_minutes,
                                },
                            }
                        )
                    )
                except sqlalchemy.exc.IntegrityError as integrity_error:
                    await db.rollback()

                    race_result = await db.execute(
                        select(ExamSession)
                        .options(noload("*"))
                        .where(
                            ExamSession.user_id == current_user.id,
                            ExamSession.exam_id == exam_id,
                            ExamSession.status.in_(("in_progress", "active")),
                        )
                        .order_by(ExamSession.start_time.desc(), ExamSession.id.desc())
                    )
                    raced_session = race_result.scalar_one_or_none()
                    if raced_session is None:
                        logger.error(
                            "EXAM_START | ACTIVE_SESSION_RACE_MISS | user=%s exam=%s error=%s",
                            current_user.id,
                            exam_id,
                            str(integrity_error),
                        )
                        raise HTTPException(
                            status_code=409,
                            detail="Konflik saat memulai sesi ujian, silakan coba lagi.",
                        )

                    is_resumed_session = True
                    session = raced_session
                    logger.warning(
                        "EXAM_START | ACTIVE_SESSION_RACE_RESUME | user=%s exam=%s session=%s",
                        current_user.id,
                        exam_id,
                        session.id,
                    )

            await db.commit()

        # Store session in Redis with idempotent timer data.
        # Do NOT overwrite started_at for an existing/resumed session.
        existing_redis_data = await get_session_data(session.id) if is_resumed_session else None
        started_at_iso = (
            (existing_redis_data or {}).get("started_at")
            or session.start_time.isoformat()
        )
        session_cache_data = {
            "session_id": session.id,
            "user_id": current_user.id,
            "exam_id": exam_id,
            "start_time": session.start_time.isoformat(),
            "end_time": session.end_time.isoformat() if session.end_time else None,
            "started_at": started_at_iso,
            "duration_seconds": exam.duration_minutes * 60,
            "elapsed_seconds": int((existing_redis_data or {}).get("elapsed_seconds") or 0),
            "paused": False,
            "duration_minutes": exam.duration_minutes,
            "status": "in_progress",
            "answered_count": int((existing_redis_data or {}).get("answered_count") or 0),
            "answered_count_stale": False,
            "total_questions": int((existing_redis_data or {}).get("total_questions") or 0),
            "violation_count": int(session.violation_count or 0),
        }
        total_paused_seconds = max(
            int((existing_redis_data or {}).get("total_paused_seconds") or 0),
            int(session.total_paused_seconds or 0),
        )
        if total_paused_seconds > 0:
            session_cache_data["total_paused_seconds"] = total_paused_seconds
        await store_session_data(session.id, session_cache_data)

        # Broadcast session start
        await _publish_exam_monitor_event(exam_id, {
            "type": "student_started",
            "user_id": current_user.id,
            "username": current_user.username,
            "session_id": session.id,
            "timestamp": now.isoformat()
        })

        questions_payload = await exam_service.get_questions_payload(exam_id)
        if not questions_payload:
            raise HTTPException(status_code=404, detail="Soal ujian tidak ditemukan")

        total_questions_from_payload = len(questions_payload)
        if int(session_cache_data.get("total_questions") or 0) != total_questions_from_payload:
            session_cache_data["total_questions"] = total_questions_from_payload
            await store_session_data(session.id, session_cache_data)

        questions = _build_start_question_responses(
            questions_payload,
            exam_id=exam.id,
            user_id=current_user.id,
            shuffle_questions=bool(exam.shuffle_questions),
            shuffle_options=bool(exam.shuffle_options),
            secret_key=settings.secret_key,
        )

        return ExamStartResponse(
            session_id=session.id,
            exam_id=exam.id,
            exam_title=exam.title,
            duration_minutes=exam.duration_minutes,
            question_count=len(questions),
            start_time=session.start_time,
            end_time=session.start_time + timedelta(minutes=exam.duration_minutes),
            server_time=datetime.now(timezone.utc),
            show_results=exam.show_results,
            show_teacher_name=exam.show_teacher_name if exam.show_teacher_name is not None else True,
            teacher_name=exam.creator.full_name if (exam.show_teacher_name and exam.creator) else None,
            subject=exam.subject,
            exam_type=exam.exam_type,
            shuffle_questions=bool(exam.shuffle_questions),
            shuffle_options=bool(exam.shuffle_options),
            session_poll_token=create_session_poll_token(
                session_id=session.id,
                user_id=current_user.id,
                expires_minutes=SESSION_POLL_TOKEN_EXPIRES_MINUTES,
            ),
            session_poll_token_expires_minutes=SESSION_POLL_TOKEN_EXPIRES_MINUTES,
            questions=questions
        )


# ============== SPRINT 1.3: NEW ENDPOINTS ==============


def _build_exam_analytics_from_aggregate_row(
    exam_id: int,
    passing_score: Optional[Decimal],
    stats: Mapping[str, Any],
) -> ExamAnalytics:
    """Build the legacy analytics response from scalar session aggregates."""
    total = int(stats["total_participants"] or 0)
    if total == 0:
        return ExamAnalytics(
            exam_id=exam_id,
            total_participants=0,
            active_sessions=0,
            completed_sessions=0,
            average_score=0,
            highest_score=0,
            lowest_score=0,
            pass_rate=0,
            score_distribution={},
            difficult_questions=[],
            violation_stats={},
        )

    completed = int(stats["completed_sessions"] or 0)
    scored_sessions = int(stats["scored_sessions"] or 0)
    passed_count = (
        int(stats["passed_sessions"] or 0)
        if passing_score
        else scored_sessions
    )
    pass_rate = (passed_count / completed * 100) if completed else 0

    return ExamAnalytics(
        exam_id=exam_id,
        total_participants=total,
        active_sessions=int(stats["active_sessions"] or 0),
        completed_sessions=completed,
        average_score=round(float(stats["average_score"] or 0), 2),
        highest_score=float(stats["highest_score"] or 0),
        lowest_score=float(stats["lowest_score"] or 0),
        pass_rate=round(pass_rate, 2),
        score_distribution={
            "0-20": int(stats["score_0_20"] or 0),
            "21-40": int(stats["score_21_40"] or 0),
            "41-60": int(stats["score_41_60"] or 0),
            "61-80": int(stats["score_61_80"] or 0),
            "81-100": int(stats["score_81_100"] or 0),
        },
        difficult_questions=[],  # Requires deeper query on Answer table
        violation_stats={"total_violations": int(stats["total_violations"] or 0)},
    )


@router.get("/{exam_id}/analytics", response_model=ExamAnalytics)
async def get_exam_analytics(
    exam_id: int,
    current_user: User = Depends(get_current_teacher),
    db: AsyncSession = Depends(get_db_read),
):
    """Get comprehensive analytics for an exam without loading session graphs."""
    exam_result = await db.execute(
        select(
            Exam.creator_id.label("creator_id"),
            Exam.passing_score.label("passing_score"),
        ).where(Exam.id == exam_id)
    )
    exam_data = exam_result.mappings().one_or_none()

    if exam_data is None:
        raise HTTPException(status_code=404, detail="Ujian tidak ditemukan")

    await _enforce_exam_owner_or_admin_access(
        db,
        current_user,
        int(exam_data["creator_id"]),
    )

    completed_statuses = ("completed", "submitted")
    completed_condition = ExamSession.status.in_(completed_statuses)
    scored_condition = and_(
        completed_condition,
        ExamSession.score.is_not(None),
    )
    passing_score = exam_data["passing_score"]
    passing_threshold = passing_score or 0

    stats_result = await db.execute(
        select(
            func.count(ExamSession.id).label("total_participants"),
            func.count(ExamSession.id)
            .filter(ExamSession.status == "in_progress")
            .label("active_sessions"),
            func.count(ExamSession.id)
            .filter(completed_condition)
            .label("completed_sessions"),
            func.count(ExamSession.id)
            .filter(scored_condition)
            .label("scored_sessions"),
            func.avg(ExamSession.score)
            .filter(scored_condition)
            .label("average_score"),
            func.max(ExamSession.score)
            .filter(scored_condition)
            .label("highest_score"),
            func.min(ExamSession.score)
            .filter(scored_condition)
            .label("lowest_score"),
            func.count(ExamSession.id)
            .filter(and_(scored_condition, ExamSession.score >= passing_threshold))
            .label("passed_sessions"),
            func.count(ExamSession.id)
            .filter(and_(scored_condition, ExamSession.score <= 20))
            .label("score_0_20"),
            func.count(ExamSession.id)
            .filter(
                and_(
                    scored_condition,
                    ExamSession.score > 20,
                    ExamSession.score <= 40,
                )
            )
            .label("score_21_40"),
            func.count(ExamSession.id)
            .filter(
                and_(
                    scored_condition,
                    ExamSession.score > 40,
                    ExamSession.score <= 60,
                )
            )
            .label("score_41_60"),
            func.count(ExamSession.id)
            .filter(
                and_(
                    scored_condition,
                    ExamSession.score > 60,
                    ExamSession.score <= 80,
                )
            )
            .label("score_61_80"),
            func.count(ExamSession.id)
            .filter(and_(scored_condition, ExamSession.score > 80))
            .label("score_81_100"),
            func.coalesce(func.sum(ExamSession.violation_count), 0).label(
                "total_violations"
            ),
        ).where(ExamSession.exam_id == exam_id)
    )

    return _build_exam_analytics_from_aggregate_row(
        exam_id,
        passing_score,
        stats_result.mappings().one(),
    )


@router.get("/{exam_id}/preview", response_model=ExamStartResponse)
async def preview_exam(
    exam_id: int,
    simulate_student_shuffle: bool = Query(
        default=False,
        description="Simulasikan urutan acak seperti saat siswa memulai ujian"
    ),
    current_user: User = Depends(get_current_teacher),
    db: AsyncSession = Depends(get_db)
):
    """Preview exam as a teacher (no SEB check, no session recording)."""
    # Fetch complete exam
    result = await db.execute(
        select(Exam)
        .options(selectinload(Exam.questions).selectinload(Question.options))
        .where(Exam.id == exam_id)
    )
    exam = result.scalar_one_or_none()

    if not exam:
        raise HTTPException(404, "Exam not found")

    await _enforce_exam_owner_or_admin_access(
        db,
        current_user,
        exam.creator_id,
    )
    start_time = datetime.now(timezone.utc)
    questions = _build_start_question_responses(
        _questions_to_start_payload(exam.questions),
        exam_id=exam.id,
        user_id=current_user.id,
        shuffle_questions=bool(simulate_student_shuffle and exam.shuffle_questions),
        shuffle_options=bool(simulate_student_shuffle and exam.shuffle_options),
        secret_key=settings.secret_key,
    )

    return ExamStartResponse(
        session_id=0, # Dummy ID
        exam_id=exam.id,
        exam_title=(
            f"[SIMULASI SISWA] {exam.title}"
            if simulate_student_shuffle
            else f"[PREVIEW] {exam.title}"
        ),
        duration_minutes=exam.duration_minutes,
        question_count=len(questions),
        start_time=start_time,
        end_time=start_time + timedelta(minutes=exam.duration_minutes),
        server_time=datetime.now(timezone.utc),  # Server time for preview consistency
        show_results=exam.show_results if exam.show_results is not None else True,
        show_teacher_name=exam.show_teacher_name if exam.show_teacher_name is not None else True,
        teacher_name=exam.creator.full_name if exam.creator else None,
        subject=exam.subject,  # FIX: Include exam metadata for preview
        exam_type=exam.exam_type,  # FIX: Include exam type for preview
        shuffle_questions=bool(exam.shuffle_questions),
        shuffle_options=bool(exam.shuffle_options),
        questions=questions
    )



async def _publish_exam_monitor_event(exam_id: int, payload: Dict[str, Any]) -> None:
    """Backward-compatible wrapper for split routers that import this helper."""
    from app.core.exam_monitor_events import publish_exam_monitor_event

    await publish_exam_monitor_event(exam_id, payload)


@router.post("/sessions/{session_id}/force-submit", response_model=ExamSubmitResponse)
async def force_submit_session(
    session_id: int,
    current_user: User = Depends(get_current_exam_monitor),
    db: AsyncSession = Depends(get_db)
):
    """
    Force submit a specific session (Teacher/Admin only).
    Useful when a student has finished but the status is stuck in 'in_progress'.
    Calculates the score based on answers currently saved in the database.
    """
    # 1. Get session with exam and questions
    stmt = (
        select(ExamSession)
        .options(
            selectinload(ExamSession.exam)
            .selectinload(Exam.questions)
            .selectinload(Question.options),
            selectinload(ExamSession.answers),
            selectinload(ExamSession.user)
        )
        .where(ExamSession.id == session_id)
    )
    result = await db.execute(stmt)
    session = result.scalar_one_or_none()

    if not session:
        raise HTTPException(status_code=404, detail="Sesi ujian tidak ditemukan")

    await _enforce_exam_owner_or_admin_access(
        db,
        current_user,
        session.exam.creator_id,
        allow_pengawas=True,
    )

    # Idempotent behavior: if session already submitted/completed, return current result
    if session.status in ("submitted", "completed"):
        total_points = sum(float(q.points) for q in session.exam.questions)
        points_earned = sum(float(a.points_earned or 0) for a in session.answers)
        percentage = (points_earned / total_points * 100) if total_points > 0 else 0
        passed = None
        if session.exam.passing_score:
            passed = percentage >= float(session.exam.passing_score)
        return ExamSubmitResponse(
            session_id=session.id,
            status="submitted",
            score=percentage,
            total_points=total_points,
            points_earned=points_earned,
            percentage=percentage,
            passed=passed,
            message="Sesi sudah pernah dikumpulkan."
        )

    if session.status != "in_progress":
        raise HTTPException(status_code=400, detail="Sesi tidak dalam status in_progress")

    submitted_at = session.end_time or datetime.now(timezone.utc)
    finalize_result = finalize_exam_session_submission(session, submitted_at=submitted_at)

    # 5. Log Action
    log = ExamLog(
        session_id=session.id,
        event_type="FORCE_SUBMIT_BY_TEACHER",
        event_data={
            "teacher_id": current_user.id,
            "teacher_name": current_user.full_name,
            "category": "admin_decision",
            "allow_continue": False,
            "score": session.score,
            "reason": "Teacher forced submission via admin panel"
        }
    )
    db.add(log)

    breakdown_log = ExamLog(
        session_id=session.id,
        event_type="SCORE_BREAKDOWN",
        event_data={"score_breakdown": finalize_result.score_breakdown}
    )
    db.add(breakdown_log)

    await db.commit()
    await _invalidate_exam_results_cache(session.exam_id)

    # 6. Broadcast Update (best-effort)
    try:
        await _publish_exam_monitor_event(session.exam_id, {
            "type": "student_submitted",
            "user_id": session.user_id,
            "username": session.user.username,
            "session_id": session.id,
            "score": float(session.score) if session.score is not None else 0,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "is_forced": True
        })
    except Exception as exc:
        logger.warning("Failed to publish forced submission for session %s: %s", session.id, str(exc))

    # Determine if passed
    passed = None
    if session.exam.passing_score:
        passed = finalize_result.percentage >= float(session.exam.passing_score)

    return ExamSubmitResponse(
        session_id=session.id,
        status="submitted",
        score=finalize_result.percentage,
        total_points=finalize_result.total_points,
        points_earned=finalize_result.points_earned,
        percentage=finalize_result.percentage,
        passed=passed,
        message="Sesi berhasil diselesaikan secara paksa."
    )
