"""
DOCX generator for assessment analysis exports (PAN / PAP).
"""
from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List
from datetime import datetime
import re

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    DOCX_AVAILABLE = True
except Exception:  # pragma: no cover - runtime dependency guard
    Document = None  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    Inches = None  # type: ignore[assignment]
    DOCX_AVAILABLE = False


DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

BASE_DIR = Path(__file__).resolve().parents[2]
PAN_TEMPLATE_PATH = BASE_DIR / "Experimental" / "PAN" / "Contoh pan_.docx"
PAP_TEMPLATE_PATH = (
    BASE_DIR / "Experimental" / "PAP" / "Laporan_PAP_UAM_Antropologi_2026.docx"
)


class AssessmentTemplateValidationError(Exception):
    """Raised when template still contains unresolved placeholders."""


def _fmt_num(value: Any, digits: int = 2) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        number = 0.0
    return f"{number:.{digits}f}"


def _fmt_num_trimmed(value: Any, digits: int = 2) -> str:
    text = _fmt_num(value, digits)
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text or "0"


def _fmt_pct(value: Any, digits: int = 2) -> str:
    return f"{_fmt_num(value, digits)}%"


def _iter_paragraphs(doc: Any) -> Iterable[Any]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _replace_placeholders(doc: Any, replacements: Dict[str, str]) -> None:
    if not replacements:
        return
    for paragraph in _iter_paragraphs(doc):
        text_value = paragraph.text or ""
        if not text_value:
            continue
        updated = text_value
        for placeholder, value in replacements.items():
            updated = updated.replace(placeholder, value)
        if updated != text_value:
            paragraph.text = updated


def _set_cell_text(cell: Any, value: Any) -> None:
    cell.text = str(value if value is not None else "")


def _clone_cell_properties(source_cell: Any, target_cell: Any) -> None:
    source_tcpr = getattr(source_cell._tc, "tcPr", None)
    if source_tcpr is None:
        return
    target_tc = target_cell._tc
    target_tcpr = getattr(target_tc, "tcPr", None)
    if target_tcpr is not None:
        target_tc.remove(target_tcpr)
    target_tc.insert(0, deepcopy(source_tcpr))


def _remove_table_column(table: Any, col_idx: int) -> None:
    if col_idx < 0 or not table.rows:
        return

    for row in table.rows:
        row_cells = row.cells
        if col_idx >= len(row_cells):
            continue
        row._tr.remove(row_cells[col_idx]._tc)

    tbl_grid = getattr(table._tbl, "tblGrid", None)
    if tbl_grid is not None:
        grid_cols = list(tbl_grid.gridCol_lst)
        if col_idx < len(grid_cols):
            tbl_grid.remove(grid_cols[col_idx])


def _normalize_pan_grade_threshold_table(pan_grade_table: Any) -> None:
    if not pan_grade_table.rows:
        return
    header_cells = pan_grade_table.rows[0].cells
    if len(header_cells) < 4:
        return
    header_values = [str(cell.text or "").strip().lower() for cell in header_cells]
    if "keterangan" not in header_values:
        return

    # Remove "Keterangan" column and enforce 3-column header.
    remove_idx = header_values.index("keterangan")
    _remove_table_column(pan_grade_table, remove_idx)
    normalized_header_cells = pan_grade_table.rows[0].cells
    if len(normalized_header_cells) >= 3:
        _set_cell_text(normalized_header_cells[0], "Kategori")
        _set_cell_text(normalized_header_cells[1], "Rumus Batas Bawah")
        _set_cell_text(normalized_header_cells[2], "Batas Skor")


def _clear_legacy_pan_signature_paragraphs(doc: Any) -> None:
    legacy_prefixes = (
        "rokan hulu,",
        "guru pelaksana,",
        "[nama guru pelaksana]",
        "nama guru pelaksana",
    )
    for paragraph in doc.paragraphs:
        text_value = str(paragraph.text or "").strip().lower()
        if not text_value:
            continue
        if any(text_value.startswith(prefix) for prefix in legacy_prefixes):
            paragraph.text = ""


def _remove_existing_pan_signature_tables(doc: Any) -> None:
    # Guard against future template variants that already include a signature block.
    removable_tables: List[Any] = []
    for table in doc.tables:
        table_text = " ".join(
            str(cell.text or "").strip().lower()
            for row in table.rows
            for cell in row.cells
        )
        if not table_text:
            continue
        if "mengetahui" in table_text and (
            "guru mata pelajaran" in table_text
            or "guru pelaksana" in table_text
            or "kepala man 1 rokan hulu" in table_text
        ):
            removable_tables.append(table)

    for table in removable_tables:
        tbl_element = table._tbl
        parent = tbl_element.getparent()
        if parent is not None:
            parent.remove(tbl_element)


def _set_paragraph_alignment_center(paragraph: Any) -> None:
    if WD_ALIGN_PARAGRAPH is not None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _force_table_grid_borders(table: Any) -> None:
    if OxmlElement is None or qn is None:
        return

    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = OxmlElement(f"w:{edge}")
        edge_element.set(qn("w:val"), "single")
        edge_element.set(qn("w:sz"), "8")
        edge_element.set(qn("w:space"), "0")
        edge_element.set(qn("w:color"), "auto")
        borders.append(edge_element)
    tbl_pr.append(borders)


def _build_pan_signature_block(doc: Any, payload: Dict[str, Any]) -> None:
    report_date = _extract_report_date(payload)
    exam = payload.get("exam", {})
    subject = str(exam.get("subject") or "").strip()
    right_title = f"Guru Mata Pelajaran {subject}" if subject else "Guru Mata Pelajaran"

    _clear_legacy_pan_signature_paragraphs(doc)
    _remove_existing_pan_signature_tables(doc)

    signature_table = doc.add_table(rows=1, cols=2)
    try:
        signature_table.style = "Table Grid"
    except Exception:
        # Some templates do not ship with this style; force borders manually.
        pass
    _force_table_grid_borders(signature_table)

    left_cell = signature_table.rows[0].cells[0]
    right_cell = signature_table.rows[0].cells[1]

    left_p1 = left_cell.paragraphs[0]
    _set_cell_text(left_cell, "")
    _set_cell_text(right_cell, "")
    left_p1 = left_cell.paragraphs[0]
    right_p1 = right_cell.paragraphs[0]

    left_p1.text = "Mengetahui,"
    _set_paragraph_alignment_center(left_p1)
    left_p2 = left_cell.add_paragraph("Kepala MAN 1 Rokan Hulu")
    _set_paragraph_alignment_center(left_p2)
    for run in left_p2.runs:
        run.bold = True
    left_cell.add_paragraph("")
    left_cell.add_paragraph("")
    left_name = left_cell.add_paragraph()
    _set_paragraph_alignment_center(left_name)
    left_name.add_run("(")
    left_name_line = left_name.add_run(" " * 31)
    left_name_line.underline = True
    left_name.add_run(")")
    left_nip = left_cell.add_paragraph()
    _set_paragraph_alignment_center(left_nip)
    left_nip.add_run("NIP. ")
    left_nip_line = left_nip.add_run(" " * 24)
    left_nip_line.underline = True

    right_p1.text = f"Rokan Hulu, {report_date}"
    _set_paragraph_alignment_center(right_p1)
    right_p2 = right_cell.add_paragraph(right_title)
    _set_paragraph_alignment_center(right_p2)
    for run in right_p2.runs:
        run.bold = True
    right_cell.add_paragraph("")
    right_cell.add_paragraph("")
    right_name = right_cell.add_paragraph()
    _set_paragraph_alignment_center(right_name)
    right_name.add_run("(")
    right_name_line = right_name.add_run(" " * 31)
    right_name_line.underline = True
    right_name.add_run(")")
    right_nip = right_cell.add_paragraph()
    _set_paragraph_alignment_center(right_nip)
    right_nip.add_run("NIP. ")
    right_nip_line = right_nip.add_run(" " * 24)
    right_nip_line.underline = True


def _ensure_pan_participant_columns(participant_table: Any) -> int:
    target_headers = [
        "No",
        "Nama Peserta",
        "Kelas",
        "Skor Mentah",
        "Nilai PAN",
        "Nilai UM",
        "Predikat",
        "Nilai Huruf",
        "Nilai Skala 10",
        "Peringkat",
    ]
    target_count = len(target_headers)

    if not participant_table.rows:
        return 0

    # Remove legacy "Kat." column when present in already-extended layouts.
    header_values = [
        str(cell.text or "").strip().lower()
        for cell in participant_table.rows[0].cells
    ]
    legacy_kat_tokens = ("kat.", "kat", "kategori")
    for token in legacy_kat_tokens:
        if token in header_values:
            _remove_table_column(participant_table, header_values.index(token))
            break

    original_col_count = len(participant_table.rows[0].cells)
    if Inches is not None:
        while len(participant_table.rows[0].cells) < target_count:
            participant_table.add_column(Inches(0.85))

    # Ensure added columns inherit cell shading/borders from template columns.
    current_col_count = len(participant_table.rows[0].cells)
    if current_col_count > original_col_count and original_col_count > 0:
        source_index = original_col_count - 1
        for row in participant_table.rows:
            row_cells = row.cells
            if not row_cells:
                continue
            source_cell = row_cells[min(source_index, len(row_cells) - 1)]
            for col_idx in range(original_col_count, len(row_cells)):
                _clone_cell_properties(source_cell, row_cells[col_idx])

    header_cells = participant_table.rows[0].cells
    if len(header_cells) >= target_count:
        for idx, header in enumerate(target_headers):
            _set_cell_text(header_cells[idx], header)

    return len(participant_table.rows[0].cells)


def _build_pan_replacements(payload: Dict[str, Any]) -> Dict[str, str]:
    exam = payload.get("exam", {})
    stats = payload.get("stats", {})
    pan = payload.get("pan", {})
    thresholds = pan.get("thresholds", {})
    scale10 = pan.get("scale10_thresholds", {})
    report_date = _extract_report_date(payload)
    teacher_name = str(exam.get("teacher_name") or "-")
    class_count_value = int(pan.get("class_count_used") or pan.get("class_count") or 1)

    return {
        "[Isi judul ujian]": str(exam.get("title") or "-"),
        "[Isi tanggal pelaksanaan]": str(exam.get("date_text") or "-"),
        "[Isi nama guru pelaksana]": str(exam.get("teacher_name") or "-"),
        "[Isi total peserta]": str(stats.get("participant_count") or 0),
        "[Isi X̅ aktual]": _fmt_num(pan.get("mean")),
        "[Isi simpangan baku (s)]": _fmt_num(pan.get("std_dev")),
        "[Skor tertinggi]": _fmt_num(stats.get("highest")),
        "[Skor terendah]": _fmt_num(stats.get("lowest")),
        "[Rentang]": _fmt_num(pan.get("score_range")),
        "[Hasil pembulatan banyak kelas]": str(class_count_value),
        "[Interval]": _fmt_num(pan.get("interval_used", pan.get("interval"))),
        "[Hasil rata-rata aktual]": _fmt_num(pan.get("mean")),
        "[Hasil simpangan baku]": _fmt_num(pan.get("std_dev")),
        "[Tanggal penetapan laporan]": report_date,
        "[Nama guru pelaksana]": teacher_name,
        "[batas A]": _fmt_num(thresholds.get("a_min")),
        "[batas B bawah]": _fmt_num(thresholds.get("b_min")),
        "[batas B atas]": _fmt_num(thresholds.get("a_min")),
        "[batas C bawah]": _fmt_num(thresholds.get("c_min")),
        "[batas C atas]": _fmt_num(thresholds.get("b_min")),
        "[batas D bawah]": _fmt_num(thresholds.get("d_min")),
        "[batas D atas]": _fmt_num(thresholds.get("c_min")),
        "[batas E]": _fmt_num(thresholds.get("d_min")),
        "[batas 10]": _fmt_num(scale10.get("10_min")),
        "[batas 9 bawah]": _fmt_num(scale10.get("9_min")),
        "[batas 9 atas]": _fmt_num(scale10.get("10_min")),
        "[batas 8 bawah]": _fmt_num(scale10.get("8_min")),
        "[batas 8 atas]": _fmt_num(scale10.get("9_min")),
        "[batas 7 bawah]": _fmt_num(scale10.get("7_min")),
        "[batas 7 atas]": _fmt_num(scale10.get("8_min")),
        "[batas 6 bawah]": _fmt_num(scale10.get("6_min")),
        "[batas 6 atas]": _fmt_num(scale10.get("7_min")),
        "[batas 5 bawah]": _fmt_num(scale10.get("5_min")),
        "[batas 5 atas]": _fmt_num(scale10.get("6_min")),
        "[batas 4 bawah]": _fmt_num(scale10.get("4_min")),
        "[batas 4 atas]": _fmt_num(scale10.get("5_min")),
        "[batas 3 bawah]": _fmt_num(scale10.get("3_min")),
        "[batas 3 atas]": _fmt_num(scale10.get("4_min")),
        "[batas 2 bawah]": _fmt_num(scale10.get("2_min")),
        "[batas 2 atas]": _fmt_num(scale10.get("3_min")),
        "[batas 1 bawah]": _fmt_num(scale10.get("1_min")),
        "[batas 1 atas]": _fmt_num(scale10.get("2_min")),
        "[batas 0]": _fmt_num(scale10.get("0_max")),
    }


def _build_pap_replacements(payload: Dict[str, Any]) -> Dict[str, str]:
    exam = payload.get("exam", {})
    stats = payload.get("stats", {})
    pap = payload.get("pap", {})
    smi_value = _fmt_num(pap.get("smi", 100), 0)
    kkm_value = _fmt_num(pap.get("kkm"), 0)
    pass_count = int(pap.get("pass_count") or 0)
    fail_count = int(pap.get("fail_count") or 0)
    participant_count = int(stats.get("participant_count") or 0)
    pass_pct = _fmt_num(pap.get("pass_percentage"), 2)
    fail_pct_value = 0.0
    if participant_count > 0:
        fail_pct_value = (float(fail_count) / float(participant_count)) * 100.0
    fail_pct = _fmt_num(fail_pct_value, 2)
    avg_score = _fmt_num(stats.get("average"), 2)
    report_date = _extract_report_date(payload)
    teacher_name = str(exam.get("teacher_name") or "-")

    return {
        "Skor Maksimum Ideal (SMI) = ________  |  KKM = ________": (
            f"Skor Maksimum Ideal (SMI) = {smi_value}  |  KKM = {kkm_value}"
        ),
        "Berikut adalah daftar nilai seluruh peserta ujian yang dikonversi menggunakan metode Penilaian Acuan Patokan (PAP) dengan KKM = ________.": (
            "Berikut adalah daftar nilai seluruh peserta ujian yang dikonversi "
            f"menggunakan metode Penilaian Acuan Patokan (PAP) dengan KKM = {kkm_value}."
        ),
        "Berdasarkan hasil penilaian dengan metode Penilaian Acuan Patokan (PAP), dari ________ peserta ujian terdapat ________ siswa (______%) yang dinyatakan TUNTAS memenuhi KKM sebesar ________. Sebanyak ________ siswa (______%) dinyatakan TIDAK TUNTAS karena nilai PAP berada di bawah KKM.": (
            "Berdasarkan hasil penilaian dengan metode Penilaian Acuan Patokan (PAP), "
            f"dari {participant_count} peserta ujian terdapat {pass_count} siswa ({pass_pct}%) "
            f"yang dinyatakan TUNTAS memenuhi KKM sebesar {kkm_value}. "
            f"Sebanyak {fail_count} siswa ({fail_pct}%) dinyatakan TIDAK TUNTAS "
            "karena nilai PAP berada di bawah KKM."
        ),
        "1. Program Remedial: Diberikan kepada siswa yang belum tuntas (jumlah: ________) dengan fokus pada materi yang belum dikuasai.": (
            "1. Program Remedial: Diberikan kepada siswa yang belum tuntas "
            f"(jumlah: {fail_count}) dengan fokus pada materi yang belum dikuasai."
        ),
        "2. Program Pengayaan: Diberikan kepada siswa yang telah mencapai KKM (jumlah: ________) untuk memperdalam dan memperluas pemahaman materi.": (
            "2. Program Pengayaan: Diberikan kepada siswa yang telah mencapai KKM "
            f"(jumlah: {pass_count}) untuk memperdalam dan memperluas pemahaman materi."
        ),
        "3. Evaluasi Soal & Pembelajaran: Bandingkan rata-rata nilai PAP (________) terhadap KKM (________) untuk menentukan tindak lanjut perbaikan pembelajaran.": (
            "3. Evaluasi Soal & Pembelajaran: Bandingkan rata-rata nilai PAP "
            f"({avg_score}) terhadap KKM ({kkm_value}) untuk menentukan tindak "
            "lanjut perbaikan pembelajaran."
        ),
        "Rokan Hulu, ____________________": f"Rokan Hulu, {report_date}",
        "_______________________________": teacher_name,
        "KKM = ________": f"KKM = {kkm_value}",
        "KKM sebesar ________": f"KKM sebesar {kkm_value}",
        "KKM = ________.": f"KKM = {kkm_value}.",
    }


def _extract_report_date(payload: Dict[str, Any]) -> str:
    generated_at = str(payload.get("generated_at") or "").strip()
    exam = payload.get("exam", {})
    exam_date = str(exam.get("date_text") or "").strip()
    for candidate in (generated_at, exam_date):
        if not candidate:
            continue
        match = re.search(r"\d{1,2}\s+[A-Za-zÀ-ÿ]+\s+\d{4}", candidate)
        if match:
            return match.group(0)
    return datetime.now().strftime("%d %B %Y")


def _build_frequency_rows(payload: Dict[str, Any], class_rows_capacity: int) -> List[Dict[str, Any]]:
    participants = payload.get("participants", [])
    scores = []
    for row in participants:
        try:
            scores.append(float(row.get("score")))
        except (TypeError, ValueError):
            continue

    if not scores:
        return []

    min_score = min(scores)
    max_score = max(scores)
    score_range = max_score - min_score
    pan = payload.get("pan", {})
    requested_classes = int(pan.get("class_count") or class_rows_capacity or 1)
    class_count = max(1, min(requested_classes, class_rows_capacity))

    interval = float(pan.get("interval") or 0.0)
    if interval <= 0:
        interval = (score_range / float(class_count)) if class_count > 0 else 0.0
    if interval <= 0:
        interval = 1.0

    # Mean duga: midpoint kelas yang memuat rata-rata aktual.
    mean_score = float(pan.get("mean") or sum(scores) / len(scores))
    mean_idx = int((mean_score - min_score) / interval) if interval > 0 else 0
    if mean_idx < 0:
        mean_idx = 0
    if mean_idx >= class_count:
        mean_idx = class_count - 1
    md_lower = min_score + (mean_idx * interval)
    md_upper = md_lower + interval
    md = (md_lower + md_upper) / 2.0

    bins: List[Dict[str, Any]] = []
    for idx in range(class_count):
        lower = min_score + (idx * interval)
        upper = lower + interval if idx < (class_count - 1) else max_score
        if idx < (class_count - 1):
            count = sum(1 for s in scores if s >= lower and s < upper)
        else:
            count = sum(1 for s in scores if s >= lower and s <= max_score)
        midpoint = (lower + upper) / 2.0
        d_value = (midpoint - md) / interval if interval > 0 else 0.0
        fd = count * d_value
        d_square = d_value * d_value
        f_d_square = count * d_square
        bins.append(
            {
                "interval": f"{_fmt_num(lower)} - {_fmt_num(upper)}",
                "midpoint": _fmt_num(midpoint),
                "f": count,
                "d": _fmt_num_trimmed(d_value),
                "fd": _fmt_num(fd),
                "d_square": _fmt_num(d_square),
                "f_d_square": _fmt_num(f_d_square),
            }
        )

    pan["class_count_used"] = class_count
    pan["interval_used"] = interval
    payload["pan"] = pan
    return bins


def _fill_pan_template(doc: Any, payload: Dict[str, Any]) -> None:
    if len(doc.tables) < 8:
        raise AssessmentTemplateValidationError("Template PAN tidak memiliki struktur tabel minimal.")

    pan = payload.get("pan", {})
    stats = payload.get("stats", {})
    participants = payload.get("participants", [])
    frequency_rows = _build_frequency_rows(payload, class_rows_capacity=max(1, len(doc.tables[1].rows) - 2))

    # T1: Distribusi frekuensi kelas interval
    freq_table = doc.tables[1]
    class_rows_capacity = max(1, len(freq_table.rows) - 2)
    for idx in range(class_rows_capacity):
        row = freq_table.rows[idx + 1]
        data = frequency_rows[idx] if idx < len(frequency_rows) else None
        if data:
            _set_cell_text(row.cells[0], data["interval"])
            _set_cell_text(row.cells[1], data["midpoint"])
            _set_cell_text(row.cells[2], str(data["f"]))
            _set_cell_text(row.cells[3], data["d"])
            _set_cell_text(row.cells[4], data["fd"])
            _set_cell_text(row.cells[5], data["d_square"])
            _set_cell_text(row.cells[6], data["f_d_square"])
        else:
            for cell in row.cells:
                _set_cell_text(cell, "")
    total_row = freq_table.rows[-1]
    total_f = sum(int(row["f"]) for row in frequency_rows)
    total_fd = sum(float(row["fd"]) for row in frequency_rows) if frequency_rows else 0.0
    total_f_d_square = (
        sum(float(row["f_d_square"]) for row in frequency_rows) if frequency_rows else 0.0
    )
    _set_cell_text(total_row.cells[0], "Jumlah")
    _set_cell_text(total_row.cells[1], "")
    _set_cell_text(total_row.cells[2], str(total_f))
    _set_cell_text(total_row.cells[3], "")
    _set_cell_text(total_row.cells[4], _fmt_num(total_fd))
    _set_cell_text(total_row.cells[5], "")
    _set_cell_text(total_row.cells[6], _fmt_num(total_f_d_square))

    # T2: Batas nilai PAN (hapus kolom "Keterangan", jadi 3 kolom).
    pan_grade_table = doc.tables[2]
    _normalize_pan_grade_threshold_table(pan_grade_table)

    # T4: Ringkasan statistik
    summary_table = doc.tables[4]
    if len(summary_table.rows) >= 5:
        _set_cell_text(summary_table.rows[1].cells[1], _fmt_num(stats.get("average")))
        _set_cell_text(summary_table.rows[2].cells[1], _fmt_num(pan.get("std_dev")))
        _set_cell_text(summary_table.rows[3].cells[1], _fmt_num(stats.get("highest")))
        _set_cell_text(summary_table.rows[4].cells[1], _fmt_num(stats.get("lowest")))

    # T5: Distribusi nilai PAN
    dist_table = doc.tables[5]
    by_grade = {
        str(item.get("grade") or "").upper(): item
        for item in pan.get("letter_distribution", [])
    }
    for row_idx, grade in enumerate(["A", "B", "C", "D", "E"], start=1):
        if row_idx >= len(dist_table.rows):
            break
        row = dist_table.rows[row_idx]
        item = by_grade.get(grade, {})
        _set_cell_text(row.cells[2], str(int(item.get("count") or 0)))
        _set_cell_text(row.cells[3], _fmt_pct(item.get("percentage")))

    # T6: Daftar peserta PAN
    participant_table = doc.tables[6]
    participant_column_count = _ensure_pan_participant_columns(participant_table)
    required_rows = len(participants) + 1  # include header
    while len(participant_table.rows) < required_rows:
        participant_table.add_row()

    for idx, row_data in enumerate(participants, start=1):
        row = participant_table.rows[idx]
        um_predicate = str(row_data.get("um_predicate") or "-")
        um_score_raw = row_data.get("um_score")
        um_score = "-" if um_score_raw in (None, "") else str(um_score_raw)
        pan_scale10_raw = row_data.get("pan_scale10")
        pan_scale10 = "-" if pan_scale10_raw in (None, "") else str(pan_scale10_raw)
        pan_letter = str(row_data.get("pan_letter") or "-")
        if participant_column_count >= 10:
            _set_cell_text(row.cells[0], str(idx))
            _set_cell_text(row.cells[1], str(row_data.get("name") or "-"))
            _set_cell_text(row.cells[2], str(row_data.get("student_class") or "-"))
            _set_cell_text(row.cells[3], _fmt_num(row_data.get("score")))
            _set_cell_text(row.cells[4], _fmt_num(row_data.get("t_score")))
            _set_cell_text(row.cells[5], um_score)
            _set_cell_text(row.cells[6], um_predicate)
            _set_cell_text(row.cells[7], pan_letter)
            _set_cell_text(row.cells[8], pan_scale10)
            _set_cell_text(row.cells[9], str(row_data.get("rank") or idx))
        else:
            # Backward compatibility for older PAN table layouts.
            _set_cell_text(row.cells[0], str(idx))
            _set_cell_text(row.cells[1], str(row_data.get("name") or "-"))
            _set_cell_text(row.cells[2], str(row_data.get("student_class") or "-"))
            _set_cell_text(row.cells[3], _fmt_num(row_data.get("score")))
            _set_cell_text(
                row.cells[4],
                f"{pan_letter} ({um_predicate})",
            )
            _set_cell_text(
                row.cells[5],
                f"{pan_scale10} / UM {um_score}",
            )
            _set_cell_text(row.cells[6], _fmt_num(row_data.get("t_score")))
            _set_cell_text(row.cells[7], str(row_data.get("rank") or idx))

    for idx in range(len(participants) + 1, len(participant_table.rows)):
        row = participant_table.rows[idx]
        _set_cell_text(row.cells[0], str(idx))
        for col in range(1, len(row.cells)):
            _set_cell_text(row.cells[col], "")

    # T7: Ringkasan konversi PAN ke UM
    um_table = doc.tables[7]
    um_by_grade = {
        str(item.get("category") or "").upper(): item
        for item in pan.get("um_conversion_summary", [])
    }
    for row_idx, grade in enumerate(["A", "B", "C", "D", "E"], start=1):
        if row_idx >= len(um_table.rows):
            break
        row = um_table.rows[row_idx]
        item = um_by_grade.get(grade, {})
        _set_cell_text(row.cells[1], str(item.get("pan_range") or "-"))
        _set_cell_text(row.cells[2], str(int(item.get("count") or 0)))
        _set_cell_text(row.cells[3], _fmt_pct(item.get("percentage")))
        _set_cell_text(row.cells[6], str(item.get("student_names") or "-"))

    # Blok TTD PAN: format dua kolom, nama penanda tangan dikosongkan.
    _build_pan_signature_block(doc, payload)


def _fill_pap_template(doc: Any, payload: Dict[str, Any]) -> None:
    if len(doc.tables) < 5:
        raise AssessmentTemplateValidationError("Template PAP tidak memiliki struktur tabel minimal.")

    pap = payload.get("pap", {})
    stats = payload.get("stats", {})
    exam = payload.get("exam", {})
    participants = payload.get("participants", [])
    smi_value = int(round(float(pap.get("smi", 100))))
    kkm_value = int(round(float(pap.get("kkm", 70))))

    # T0: Metadata dokumen
    meta_table = doc.tables[0]
    if len(meta_table.rows) >= 5:
        _set_cell_text(meta_table.rows[0].cells[2], str(exam.get("title") or "-"))
        _set_cell_text(meta_table.rows[1].cells[2], str(exam.get("date_text") or "-"))
        _set_cell_text(meta_table.rows[2].cells[2], str(exam.get("teacher_name") or "-"))
        _set_cell_text(meta_table.rows[3].cells[2], str(int(stats.get("participant_count") or 0)))
        _set_cell_text(meta_table.rows[4].cells[2], f"{kkm_value} (Penilaian Acuan Patokan)")

    # T2: Ringkasan statistik
    summary_table = doc.tables[2]
    if len(summary_table.rows) >= 6:
        _set_cell_text(summary_table.rows[1].cells[1], _fmt_num(stats.get("average")))
        _set_cell_text(summary_table.rows[2].cells[1], _fmt_num(stats.get("highest")))
        _set_cell_text(summary_table.rows[3].cells[1], _fmt_num(stats.get("lowest")))
        _set_cell_text(summary_table.rows[4].cells[0], f"TUNTAS (>= KKM {kkm_value})")
        _set_cell_text(summary_table.rows[4].cells[1], str(int(pap.get("pass_count") or 0)))
        _set_cell_text(summary_table.rows[5].cells[1], str(int(pap.get("fail_count") or 0)))

    # T3: Distribusi huruf PAP
    distribution_table = doc.tables[3]
    grade_map = {
        str(item.get("grade") or "").upper(): item
        for item in pap.get("grade_distribution", [])
    }
    for row_idx, grade in enumerate(["A", "B", "C", "D", "E"], start=1):
        if row_idx >= len(distribution_table.rows):
            break
        row = distribution_table.rows[row_idx]
        item = grade_map.get(grade, {})
        _set_cell_text(row.cells[2], str(int(item.get("count") or 0)))
        _set_cell_text(row.cells[3], _fmt_pct(item.get("percentage")))

    # T4: Daftar peserta PAP
    participant_table = doc.tables[4]
    required_rows = len(participants) + 1
    while len(participant_table.rows) < required_rows:
        participant_table.add_row()
    for idx, row_data in enumerate(participants, start=1):
        row = participant_table.rows[idx]
        score_text = _fmt_num(row_data.get("score"))
        _set_cell_text(row.cells[0], str(idx))
        _set_cell_text(row.cells[1], str(row_data.get("name") or "-"))
        _set_cell_text(row.cells[2], str(row_data.get("student_class") or "-"))
        _set_cell_text(row.cells[3], score_text)
        _set_cell_text(row.cells[4], score_text)
        _set_cell_text(row.cells[5], str(row_data.get("pap_letter") or "-"))
        _set_cell_text(row.cells[6], str(row_data.get("pap_status") or "-"))
    for idx in range(len(participants) + 1, len(participant_table.rows)):
        row = participant_table.rows[idx]
        _set_cell_text(row.cells[0], str(idx))
        for col in range(1, len(row.cells)):
            _set_cell_text(row.cells[col], "")

    # Update paragraph strings with analytic conclusions + signature.
    replacements = _build_pap_replacements(payload)
    replacements.setdefault(
        "Skor Maksimum Ideal (SMI) = ________  |  KKM = ________",
        f"Skor Maksimum Ideal (SMI) = {smi_value}  |  KKM = {kkm_value}",
    )
    _replace_placeholders(doc, replacements)


def _ensure_no_unresolved_placeholders(doc: Any, report_model: str) -> None:
    unresolved_pattern = re.compile(r"\[[^\]]+\]|_{3,}")
    allowed_square_bracket_tokens = {"[(X – μ) / σ]", "[(X - μ) / σ]"}
    findings: List[str] = []

    def _is_unresolved(value: str) -> bool:
        normalized = value
        for token in allowed_square_bracket_tokens:
            normalized = normalized.replace(token, "")
        return bool(unresolved_pattern.search(normalized))

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text_value = (paragraph.text or "").strip()
        if text_value and _is_unresolved(text_value):
            findings.append(f"P{paragraph_index}:{text_value[:120]}")
        if len(findings) >= 12:
            break

    if len(findings) < 12:
        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    text_value = (cell.text or "").strip()
                    if text_value and _is_unresolved(text_value):
                        findings.append(
                            f"T{table_index}R{row_index}C{col_index}:{text_value[:100]}"
                        )
                    if len(findings) >= 12:
                        break
                if len(findings) >= 12:
                    break
            if len(findings) >= 12:
                break

    if findings:
        joined = "; ".join(findings)
        raise AssessmentTemplateValidationError(
            f"Template {report_model.upper()} masih memiliki placeholder: {joined}"
        )


def _prepare_pan_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    prepared = dict(payload)
    prepared["pan"] = dict(payload.get("pan", {}))
    prepared["stats"] = dict(payload.get("stats", {}))
    prepared["exam"] = dict(payload.get("exam", {}))
    return prepared


def _prepare_pap_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    prepared = dict(payload)
    prepared["pap"] = dict(payload.get("pap", {}))
    prepared["stats"] = dict(payload.get("stats", {}))
    prepared["exam"] = dict(payload.get("exam", {}))
    prepared["pap"].setdefault("smi", 100)
    return prepared


def generate_assessment_docx(report_model: str, payload: Dict[str, Any]) -> bytes:
    """
    Generate a DOCX file for assessment analysis export.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")

    normalized_model = str(report_model or "").strip().lower()
    if normalized_model not in {"pan", "pap"}:
        raise ValueError("Unsupported report model")

    template_path = PAN_TEMPLATE_PATH if normalized_model == "pan" else PAP_TEMPLATE_PATH
    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")

    document = Document(str(template_path))
    if normalized_model == "pan":
        prepared_payload = _prepare_pan_payload(payload)
        _replace_placeholders(document, _build_pan_replacements(prepared_payload))
        _fill_pan_template(document, prepared_payload)
        _replace_placeholders(document, _build_pan_replacements(prepared_payload))
    else:
        prepared_payload = _prepare_pap_payload(payload)
        _fill_pap_template(document, prepared_payload)
        _replace_placeholders(document, _build_pap_replacements(prepared_payload))

    _ensure_no_unresolved_placeholders(document, normalized_model)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
