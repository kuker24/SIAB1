#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
=============================================================================
FULL CODEBASE DOCUMENTATION GENERATOR - ENHANCED VERSION
=============================================================================
Script untuk menghasilkan dokumentasi DOCX lengkap dari seluruh codebase.

ENHANCED FEATURES (v4.0):
- Table of Contents (Daftar Isi)
- Overview Project + README
- API Endpoints Registry (FastAPI routes)
- Database Models Documentation (SQLAlchemy)
- Configuration/Environment Variables
- Dependencies List (requirements.txt)
- Struktur Direktori Visual (Tree)
- Statistik Codebase (LOC, file count, dll)
- Dokumentasi Detail per File dengan syntax info
- Ekstraksi class dan function dari Python files
- JavaScript Function Extraction
- HTML Template Analysis
- Professional formatting
- Grading, Analytics, Monitoring API support
- Matching Question System Documentation
- Real-time Monitoring Features
- Backup and Restore System
- APK Builder Integration
- SEB Configuration System
- WebSocket Support
- Results Display System
- Executive Summary Section
- Professional Header/Footer Layout
- CLI Arguments for Custom Output
- Optional Compact Mode (without full source code)
- Readability Controls (tree depth and code line limits)

Author: Auto-Generated Documentation Tool
Version: 4.0 (Professional Documentation Edition - March 2026)
Date: 2026-03-10
=============================================================================
"""


import os
import sys
import re
import ast
import argparse
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from collections import defaultdict

# ============================================================================
# DEPENDENCY CHECK
# ============================================================================
def ensure_docx_dependency():
    """Ensure required dependency is available."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        return True
    except ImportError as exc:
        print(
            "[ERROR] Dependency 'python-docx' belum terpasang.\n"
            "Install dulu dengan:\n"
            "  pip install python-docx",
            file=sys.stderr,
        )
        raise SystemExit(1) from exc

# Fix Windows console encoding
import sys
import io
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

Document = None
Pt = None
Inches = None
RGBColor = None
Cm = None
WD_PARAGRAPH_ALIGNMENT = None
WD_STYLE_TYPE = None
WD_TABLE_ALIGNMENT = None
qn = None
OxmlElement = None


def load_docx_symbols(auto_install=False):
    """Load python-docx symbols lazily so --help can run without docx import side-effects."""
    global Document, Pt, Inches, RGBColor, Cm
    global WD_PARAGRAPH_ALIGNMENT, WD_STYLE_TYPE, WD_TABLE_ALIGNMENT
    global qn, OxmlElement

    try:
        from docx import Document as _Document
        from docx.shared import Pt as _Pt, Inches as _Inches, RGBColor as _RGBColor, Cm as _Cm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _WD_PARAGRAPH_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE as _WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement
    except ImportError:
        if not auto_install:
            ensure_docx_dependency()
            raise SystemExit(1)

        import subprocess
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

        from docx import Document as _Document
        from docx.shared import Pt as _Pt, Inches as _Inches, RGBColor as _RGBColor, Cm as _Cm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as _WD_PARAGRAPH_ALIGNMENT
        from docx.enum.style import WD_STYLE_TYPE as _WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement

    Document = _Document
    Pt = _Pt
    Inches = _Inches
    RGBColor = _RGBColor
    Cm = _Cm
    WD_PARAGRAPH_ALIGNMENT = _WD_PARAGRAPH_ALIGNMENT
    WD_STYLE_TYPE = _WD_STYLE_TYPE
    WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
    qn = _qn
    OxmlElement = _OxmlElement

# ============================================================================
# XML SANITIZATION
# ============================================================================
def sanitize_for_xml(text):
    """
    Remove control characters that are not allowed in XML.
    XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
    """
    if not text:
        return text

    # Remove NULL bytes and other control characters (except tab, newline, carriage return)
    cleaned = []
    for char in text:
        code = ord(char)
        # Allow: tab (9), newline (10), carriage return (13), and printable chars (32+)
        if code == 9 or code == 10 or code == 13 or (code >= 32 and code != 127):
            cleaned.append(char)
        else:
            # Replace control characters with space
            cleaned.append(' ')

    return ''.join(cleaned)

# ============================================================================
# CONFIGURATION
# ============================================================================
class Config:
    """Configuration for documentation generator."""

    # Directories to exclude
    EXCLUDE_DIRS = {
        '__pycache__', '.git', '.idea', 'venv', 'env', 'node_modules',
        'build', 'dist', '.vscode', '.gemini', 'tmp', '.cache',
        'migrations', '.pytest_cache', '.mypy_cache', 'eggs', '*.egg-info',
        'apk_builds', 'flutter_client', '.dart_tool', '.android', '.gradle',
        '.venv', '.venv-docgen', '.ruff_cache'
    }

    # File extensions to exclude (binary/non-text)
    EXCLUDE_EXTENSIONS = {
        '.pyc', '.pyo', '.pyd', '.dll', '.so', '.exe', '.db', '.sqlite',
        '.sqlite3', '.log', '.zip', '.tar', '.gz', '.rar', '.7z',
        '.png', '.jpg', '.jpeg', '.gif', '.ico', '.pdf', '.docx', '.xlsx',
        '.bin', '.seb', '.lock', '.woff', '.woff2', '.ttf', '.eot', '.map',
        '.apk', '.aab', '.ipa', '.dex', '.class', '.jar', '.aar'
    }

    # Specific files to exclude
    EXCLUDE_FILES = {
        '.env', 'package-lock.json', 'yarn.lock', 'poetry.lock',
        '.DS_Store', 'Thumbs.db', '.gitignore', '.dockerignore'
    }

    # Language mapping for syntax documentation
    LANGUAGE_MAP = {
        '.py': 'Python',
        '.js': 'JavaScript',
        '.ts': 'TypeScript',
        '.tsx': 'TypeScript (React)',
        '.jsx': 'JavaScript (React)',
        '.html': 'HTML',
        '.css': 'CSS',
        '.scss': 'SCSS',
        '.sass': 'SASS',
        '.json': 'JSON',
        '.yaml': 'YAML',
        '.yml': 'YAML',
        '.md': 'Markdown',
        '.sql': 'SQL',
        '.sh': 'Shell/Bash',
        '.bat': 'Batch',
        '.ps1': 'PowerShell',
        '.dart': 'Dart/Flutter',
        '.kt': 'Kotlin',
        '.java': 'Java',
        '.xml': 'XML',
        '.dockerfile': 'Dockerfile',
        '.toml': 'TOML',
        '.ini': 'INI',
        '.cfg': 'Config',
        '.conf': 'Config',
        '.env.example': 'Environment Config',
        '.txt': 'Plain Text',
    }


@dataclass
class GeneratorOptions:
    """Runtime options for generator behavior and document identity."""

    document_title: str = "DOKUMENTASI LENGKAP CODEBASE"
    document_subtitle: str = "Laporan Teknis Otomatis"
    author_name: str = "Auto-Generated Documentation Tool"
    organization: str = ""
    include_source_code: bool = True
    max_code_lines_per_file: int = 5000
    max_tree_depth: int = 6


# ============================================================================
# DOCUMENT STYLING HELPERS
# ============================================================================
class DocStyles:
    """Helper class for document styling."""

    @staticmethod
    def set_cell_shading(cell, fill_color):
        """Set background color for table cell."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def add_horizontal_line(doc):
        """Add a horizontal line to the document."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("─" * 80)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.size = Pt(8)

    @staticmethod
    def create_styles(doc):
        """Create custom styles for the document."""
        styles = doc.styles

        # Code style
        if 'Code' not in [s.name for s in styles]:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(9)
            code_style.paragraph_format.space_before = Pt(3)
            code_style.paragraph_format.space_after = Pt(3)

        # File Path style
        if 'FilePath' not in [s.name for s in styles]:
            path_style = styles.add_style('FilePath', WD_STYLE_TYPE.PARAGRAPH)
            path_style.font.name = 'Consolas'
            path_style.font.size = Pt(10)
            path_style.font.bold = True
            path_style.font.color.rgb = RGBColor(0, 102, 204)


# ============================================================================
# PYTHON CODE ANALYZER (ENHANCED)
# ============================================================================
class PythonAnalyzer:
    """Analyze Python files to extract classes, functions, and API routes."""

    @staticmethod
    def analyze_file(filepath):
        """Extract classes, functions, and docstrings from a Python file."""
        result = {
            'classes': [],
            'functions': [],
            'imports': [],
            'module_docstring': None,
            'api_routes': [],      # NEW: FastAPI routes
            'models': [],          # NEW: SQLAlchemy models
            'decorators': []       # NEW: Decorated functions
        }

        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()

            tree = ast.parse(content)

            # Module docstring
            result['module_docstring'] = ast.get_docstring(tree)

            # Imports
            for node in ast.walk(tree):
                if isinstance(node, ast.Import):
                    for alias in node.names:
                        result['imports'].append(alias.name)
                elif isinstance(node, ast.ImportFrom):
                    module = node.module or ''
                    for alias in node.names:
                        result['imports'].append(f"{module}.{alias.name}")

            # Classes and Functions at module level
            for node in ast.iter_child_nodes(tree):
                if isinstance(node, ast.ClassDef):
                    class_info = {
                        'name': node.name,
                        'docstring': ast.get_docstring(node),
                        'methods': [],
                        'line': node.lineno,
                        'bases': [PythonAnalyzer._get_name(base) for base in node.bases]
                    }

                    # Check if it's a SQLAlchemy model
                    if any('Base' in str(base) or 'Model' in str(base) for base in class_info['bases']):
                        result['models'].append(class_info)

                    # Get methods
                    for item in node.body:
                        if isinstance(item, ast.FunctionDef):
                            method_info = {
                                'name': item.name,
                                'docstring': ast.get_docstring(item),
                                'args': [arg.arg for arg in item.args.args],
                                'line': item.lineno
                            }
                            class_info['methods'].append(method_info)
                    result['classes'].append(class_info)

                elif isinstance(node, ast.FunctionDef):
                    func_info = {
                        'name': node.name,
                        'docstring': ast.get_docstring(node),
                        'args': [arg.arg for arg in node.args.args],
                        'line': node.lineno,
                        'decorators': []
                    }

                    # Extract decorators (for API routes)
                    for decorator in node.decorator_list:
                        dec_str = PythonAnalyzer._get_decorator_string(decorator)
                        func_info['decorators'].append(dec_str)

                        # Check for API routes
                        if any(method in dec_str for method in ['.get', '.post', '.put', '.delete', '.patch']):
                            route_info = PythonAnalyzer._extract_route_info(decorator, func_info)
                            if route_info:
                                result['api_routes'].append(route_info)

                    result['functions'].append(func_info)

            return result

        except Exception as e:
            return {'error': str(e)}

    @staticmethod
    def _get_name(node):
        """Get name from AST node."""
        if isinstance(node, ast.Name):
            return node.id
        elif isinstance(node, ast.Attribute):
            return f"{PythonAnalyzer._get_name(node.value)}.{node.attr}"
        return str(node)

    @staticmethod
    def _get_decorator_string(decorator):
        """Convert decorator AST to string."""
        if isinstance(decorator, ast.Call):
            func = decorator.func
            if isinstance(func, ast.Attribute):
                return f"@{PythonAnalyzer._get_name(func.value)}.{func.attr}"
            elif isinstance(func, ast.Name):
                return f"@{func.id}"
        elif isinstance(decorator, ast.Attribute):
            return f"@{PythonAnalyzer._get_name(decorator.value)}.{decorator.attr}"
        elif isinstance(decorator, ast.Name):
            return f"@{decorator.id}"
        return "@unknown"

    @staticmethod
    def _extract_route_info(decorator, func_info):
        """Extract API route information from decorator."""
        if isinstance(decorator, ast.Call):
            func = decorator.func
            if isinstance(func, ast.Attribute):
                method = func.attr.upper()
                # Check if it's an HTTP method
                if method in ['GET', 'POST', 'PUT', 'DELETE', 'PATCH', 'OPTIONS', 'HEAD']:
                    path = ""
                    if decorator.args:
                        first_arg = decorator.args[0]
                        if isinstance(first_arg, ast.Constant):
                            path = first_arg.value
                        elif isinstance(first_arg, ast.Str):  # Python 3.7 compat
                            path = first_arg.s
                    return {
                        'method': method,
                        'path': path,
                        'function': func_info['name'],
                        'docstring': func_info['docstring']
                    }
        return None


# ============================================================================
# JAVASCRIPT ANALYZER
# ============================================================================
class JavaScriptAnalyzer:
    """Analyze JavaScript files to extract functions and classes."""

    @staticmethod
    def analyze_file(filepath):
        """Extract functions and classes from JavaScript file."""
        result = {
            'functions': [],
            'classes': [],
            'exports': []
        }

        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()

            # Extract function declarations
            func_pattern = r'(?:async\s+)?function\s+(\w+)\s*\([^)]*\)'
            for match in re.finditer(func_pattern, content):
                result['functions'].append({
                    'name': match.group(1),
                    'type': 'declaration'
                })

            # Extract arrow functions assigned to const/let/var
            arrow_pattern = r'(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s+)?\([^)]*\)\s*=>'
            for match in re.finditer(arrow_pattern, content):
                result['functions'].append({
                    'name': match.group(1),
                    'type': 'arrow'
                })

            # Extract classes
            class_pattern = r'class\s+(\w+)(?:\s+extends\s+(\w+))?'
            for match in re.finditer(class_pattern, content):
                result['classes'].append({
                    'name': match.group(1),
                    'extends': match.group(2)
                })

            # Extract exports
            export_pattern = r'export\s+(?:default\s+)?(?:const|let|var|function|class)\s+(\w+)'
            for match in re.finditer(export_pattern, content):
                result['exports'].append(match.group(1))

            return result

        except Exception as e:
            return {'error': str(e)}


# ============================================================================
# HTML ANALYZER
# ============================================================================
class HTMLAnalyzer:
    """Analyze HTML files to extract structure."""

    @staticmethod
    def analyze_file(filepath):
        """Extract structure from HTML file."""
        result = {
            'title': None,
            'scripts': [],
            'styles': [],
            'forms': [],
            'ids': [],
            'classes': set()
        }

        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()

            # Extract title
            title_match = re.search(r'<title>([^<]+)</title>', content, re.I)
            if title_match:
                result['title'] = title_match.group(1).strip()

            # Extract scripts
            script_pattern = r'<script[^>]*src=["\']([^"\']+)["\']'
            result['scripts'] = re.findall(script_pattern, content)

            # Extract stylesheets
            style_pattern = r'<link[^>]*href=["\']([^"\']+\.css[^"\']*)["\']'
            result['styles'] = re.findall(style_pattern, content)

            # Extract form actions
            form_pattern = r'<form[^>]*(?:action=["\']([^"\']+)["\'])?[^>]*(?:method=["\']([^"\']+)["\'])?'
            for match in re.finditer(form_pattern, content, re.I):
                result['forms'].append({
                    'action': match.group(1) or '#',
                    'method': match.group(2) or 'GET'
                })

            # Extract IDs
            id_pattern = r'id=["\']([^"\']+)["\']'
            result['ids'] = re.findall(id_pattern, content)

            return result

        except Exception as e:
            return result

        except Exception as e:
            return {'error': str(e)}


# ============================================================================
# DART ANALYZER (NEW)
# ============================================================================
class DartAnalyzer:
    """Analyze Dart/Flutter files to extract classes and functions."""

    @staticmethod
    def analyze_file(filepath):
        """Extract structure from Dart file."""
        result = {
            'classes': [],
            'functions': [],
            'imports': []
        }

        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()

            # Extract imports
            import_pattern = r'import\s+[\'"]([^\'"]+)[\'"]'
            result['imports'] = re.findall(import_pattern, content)

            # Extract classes
            # class ClassName extends BaseClass {
            class_pattern = r'class\s+(\w+)(?:\s+extends\s+(\w+))?(?:\s+implements\s+(\w+))?'
            for match in re.finditer(class_pattern, content):
                result['classes'].append({
                    'name': match.group(1),
                    'extends': match.group(2),
                    'implements': match.group(3)
                })

            # Extract basic functions (void main, etc.)
            func_pattern = r'(?:void|Future<[^>]+>|Widget|String|int|bool|double)\s+(\w+)\s*\('
            for match in re.finditer(func_pattern, content):
                name = match.group(1)
                # Filter out common reserved words if matched by mistake
                if name not in ['if', 'while', 'for', 'switch', 'catch']:
                    result['functions'].append(name)

            return result

        except Exception as e:
            return {'error': str(e)}


# ============================================================================
# DEPENDENCIES ANALYZER
# ============================================================================
class DependenciesAnalyzer:
    """Analyze project dependencies."""

    @staticmethod
    def parse_requirements(filepath):
        """Parse requirements.txt file."""
        dependencies = []

        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        # Parse package name and version
                        match = re.match(r'([a-zA-Z0-9_-]+)(?:[=<>!]+(.+))?', line)
                        if match:
                            dependencies.append({
                                'name': match.group(1),
                                'version': match.group(2) or 'latest',
                                'raw': line
                            })
        except Exception as e:
            return [{'error': str(e)}]

        return dependencies

    @staticmethod
    def parse_env_example(filepath):
        """Parse .env.example file."""
        env_vars = []

        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        match = re.match(r'([A-Z_][A-Z0-9_]*)=(.*)$', line)
                        if match:
                            env_vars.append({
                                'name': match.group(1),
                                'example_value': match.group(2),
                                'description': ''
                            })
                    elif line.startswith('#'):
                        # This might be a description for the next variable
                        pass
        except Exception as e:
            return [{'error': str(e)}]

        return env_vars


# ============================================================================
# DIRECTORY TREE GENERATOR
# ============================================================================
class DirectoryTree:
    """Generate a visual directory tree."""

    @staticmethod
    def generate(root_dir, exclude_dirs=None, max_depth=6):
        """Generate directory tree as a list of strings."""
        if exclude_dirs is None:
            exclude_dirs = Config.EXCLUDE_DIRS

        tree_lines = []

        def _tree(dir_path, prefix="", depth=0):
            if depth > max_depth:
                return

            try:
                entries = sorted(os.listdir(dir_path))
            except PermissionError:
                return

            # Filter out excluded directories
            entries = [e for e in entries if e not in exclude_dirs]

            dirs = [e for e in entries if os.path.isdir(os.path.join(dir_path, e))]
            files = [e for e in entries if os.path.isfile(os.path.join(dir_path, e))]

            # Sort: directories first, then files
            all_entries = dirs + files

            for i, entry in enumerate(all_entries):
                is_last = (i == len(all_entries) - 1)
                connector = "└── " if is_last else "├── "
                entry_path = os.path.join(dir_path, entry)

                if os.path.isdir(entry_path):
                    tree_lines.append(f"{prefix}{connector}📁 {entry}/")
                    extension = "    " if is_last else "│   "
                    _tree(entry_path, prefix + extension, depth + 1)
                else:
                    # Add file with appropriate icon
                    icon = DirectoryTree._get_file_icon(entry)
                    tree_lines.append(f"{prefix}{connector}{icon} {entry}")

        root_name = os.path.basename(root_dir) or root_dir
        tree_lines.append(f"📁 {root_name}/")
        _tree(root_dir)

        return tree_lines

    @staticmethod
    def _get_file_icon(filename):
        """Get appropriate icon for file type."""
        ext = os.path.splitext(filename)[1].lower()
        icons = {
            '.py': '🐍',
            '.js': '📜',
            '.ts': '📘',
            '.html': '🌐',
            '.css': '🎨',
            '.json': '📋',
            '.md': '📄',
            '.sql': '🗃️',
            '.sh': '⚙️',
            '.yaml': '📝',
            '.yml': '📝',
            '.dart': '🎯',
            '.dockerfile': '🐳',
        }
        return icons.get(ext, '📄')


# ============================================================================
# CODEBASE STATISTICS
# ============================================================================
class CodebaseStats:
    """Calculate codebase statistics."""

    def __init__(self):
        self.total_files = 0
        self.total_lines = 0
        self.total_bytes = 0
        self.files_by_type = defaultdict(int)
        self.lines_by_type = defaultdict(int)
        self.file_list = []
        self.api_routes = []
        self.models = []
        self.dependencies = []
        self.env_vars = []

    def add_file(self, filepath, content, ext):
        """Add file to statistics."""
        self.total_files += 1
        lines = len(content.splitlines()) if isinstance(content, str) else len(content)
        self.total_lines += lines
        self.total_bytes += len(content.encode('utf-8')) if isinstance(content, str) else 0 # Approx if raw lines
        self.files_by_type[ext] += 1
        self.lines_by_type[ext] += lines
        self.file_list.append({
            'path': filepath,
            'lines': lines,
            'size': len(content.encode('utf-8')) if isinstance(content, str) else 0,
            'type': ext
        })

    def add_api_routes(self, routes):
        """Add API routes to registry."""
        self.api_routes.extend(routes)

    def add_models(self, models):
        """Add database models to registry."""
        self.models.extend(models)

    def get_summary(self):
        """Get statistics summary."""
        return {
            'total_files': self.total_files,
            'total_lines': self.total_lines,
            'total_bytes': self.total_bytes,
            'files_by_type': dict(self.files_by_type),
            'lines_by_type': dict(self.lines_by_type),
            'largest_files': sorted(self.file_list, key=lambda x: x['lines'], reverse=True)[:10],
            'api_routes': self.api_routes,
            'models': self.models,
            'dependencies': self.dependencies,
            'env_vars': self.env_vars
        }


# ============================================================================
# MAIN DOCUMENTATION GENERATOR
# ============================================================================
class FullDocumentationGenerator:
    """Main class to generate comprehensive codebase documentation."""

    def __init__(self, root_dir, output_file="FULL_CODEBASE_DOCUMENTATION.docx", options=None):
        self.root_dir = os.path.abspath(root_dir)
        self.output_file = output_file
        self.doc = Document()
        self.stats = CodebaseStats()
        self.toc_entries = []
        self.options = options or GeneratorOptions()

    def generate(self):
        """Generate the full documentation."""
        print("=" * 60)
        print("FULL CODEBASE DOCUMENTATION GENERATOR v4.0")
        print("=" * 60)
        print(f"Root Directory: {self.root_dir}")
        print(f"Output File: {self.output_file}")
        print(f"Include Source Code: {'yes' if self.options.include_source_code else 'no'}")
        print("=" * 60)

        # Configure page, fonts, and header/footer
        self._configure_document_layout()

        # Create custom styles
        DocStyles.create_styles(self.doc)

        # Pre-scan for dependencies and env
        self._prescan_project()

        # 1. ERROR FIX: Perform initial scan to calculate stats BEFORE generating document
        # This allows us to put accurate "Total Lines" on the cover page and stats section
        self._perform_initial_scan()

        # Add sections
        self._add_cover_page()
        self._add_table_of_contents()
        self._add_executive_summary()
        self._add_project_overview()
        self._add_api_registry()
        self._add_database_models()
        self._add_dependencies_section()
        self._add_environment_variables()
        self._add_directory_structure()
        self._add_statistics_section()
        self._add_all_files_documentation()
        self._add_appendix()

        # Save document
        output_path = os.path.join(self.root_dir, self.output_file)
        self.doc.save(output_path)

        print("\n" + "=" * 60)
        print("DOCUMENTATION GENERATED SUCCESSFULLY!")
        print("=" * 60)
        print(f"✓ Total Files Documented: {self.stats.total_files}")
        print(f"✓ Total Lines of Code: {self.stats.total_lines:,}")
        print(f"✓ Total Size: {self.stats.total_bytes / 1024:.2f} KB")
        print(f"✓ API Endpoints Found: {len(self.stats.api_routes)}")
        print(f"✓ Database Models Found: {len(self.stats.models)}")
        print(f"✓ Dependencies: {len(self.stats.dependencies)}")
        print(f"✓ Saved to: {output_path}")
        print("=" * 60)

        return output_path

    def _prescan_project(self):
        """Pre-scan project for dependencies and environment variables."""
        print("📡 Pre-scanning project...")

        # Parse requirements.txt
        req_path = os.path.join(self.root_dir, 'requirements.txt')
        if os.path.exists(req_path):
            self.stats.dependencies = DependenciesAnalyzer.parse_requirements(req_path)
            print(f"  ✓ Found {len(self.stats.dependencies)} dependencies")

        # Parse .env.example
        env_path = os.path.join(self.root_dir, '.env.example')
        if os.path.exists(env_path):
            self.stats.env_vars = DependenciesAnalyzer.parse_env_example(env_path)
            print(f"  ✓ Found {len(self.stats.env_vars)} environment variables")

    def _perform_initial_scan(self):
        """Scan all files to calculate statistics before document generation."""
        print("🔍 Performing initial code scan...")

        for dirpath, dirnames, filenames in os.walk(self.root_dir):
            # Skip excluded directories
            dirnames[:] = [d for d in dirnames if d not in Config.EXCLUDE_DIRS]

            for filename in filenames:
                filepath = os.path.join(dirpath, filename)
                rel_path = os.path.relpath(filepath, self.root_dir)

                # Get extension
                _, ext = os.path.splitext(filename)
                ext = ext.lower()

                # Skip excluded files
                if ext in Config.EXCLUDE_EXTENSIONS or filename in Config.EXCLUDE_FILES:
                    continue

                # Skip the output file itself
                if filename == self.output_file:
                    continue

                try:
                    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                        content = f.read()
                        self.stats.add_file(rel_path, content, ext)
                except Exception:
                    pass # Skip errors in scan

        print(f"  ✓ Initial scan complete: {self.stats.total_lines} lines in {self.stats.total_files} files")

    def _configure_document_layout(self):
        """Apply professional page layout, typography, and header/footer."""
        for section in self.doc.sections:
            section.top_margin = Cm(2.2)
            section.bottom_margin = Cm(2.2)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

            # Header
            header = section.header.paragraphs[0]
            header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            header.text = f"{self.options.document_title} | {os.path.basename(self.root_dir)}"
            if header.runs:
                header.runs[0].font.size = Pt(8)
                header.runs[0].font.color.rgb = RGBColor(120, 120, 120)

            # Footer
            footer = section.footer.paragraphs[0]
            footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            footer.text = f"Dibuat otomatis pada {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            if footer.runs:
                footer.runs[0].font.size = Pt(8)
                footer.runs[0].font.color.rgb = RGBColor(120, 120, 120)

        # Base text style
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(10.5)

    def _add_executive_summary(self):
        """Add an executive summary section for easy reading."""
        stats = self.stats.get_summary()
        self.doc.add_heading('1. 📌 EXECUTIVE SUMMARY', level=1)

        self.doc.add_paragraph(
            "Bagian ini memberikan ringkasan singkat agar pembaca non-teknis dapat memahami "
            "kondisi codebase tanpa harus membaca seluruh detail file."
        )

        summary_table = self.doc.add_table(rows=1, cols=2)
        summary_table.style = 'Table Grid'
        summary_points = [
            ("Total file terdeteksi", str(stats['total_files'])),
            ("Total baris kode", f"{stats['total_lines']:,}"),
            ("Total ukuran source", f"{stats['total_bytes'] / 1024:.2f} KB"),
            ("Jumlah endpoint API", str(len(stats['api_routes']))),
            ("Jumlah model database", str(len(stats['models']))),
            ("Jumlah dependency", str(len(stats['dependencies']))),
        ]

        for key, value in summary_points:
            row = summary_table.add_row().cells
            row[0].text = key
            row[1].text = value
            row[0].paragraphs[0].runs[0].font.bold = True

        # remove initial empty row
        summary_table._tbl.remove(summary_table.rows[0]._tr)

        self.doc.add_paragraph()
        self.doc.add_heading("Temuan Utama", level=2)

        top_langs = sorted(
            stats['lines_by_type'].items(),
            key=lambda item: item[1],
            reverse=True
        )[:3]

        for ext, lines in top_langs:
            lang = Config.LANGUAGE_MAP.get(ext, ext or "(no extension)")
            pct = (lines / stats['total_lines'] * 100) if stats['total_lines'] else 0
            self.doc.add_paragraph(f"• {lang}: {lines:,} baris ({pct:.1f}%)")

        self.doc.add_heading("Rekomendasi Baca Dokumen", level=2)
        self.doc.add_paragraph("1. Mulai dari Project Overview untuk memahami konteks aplikasi.")
        self.doc.add_paragraph("2. Lanjut ke API Registry dan Database Models untuk arsitektur backend.")
        self.doc.add_paragraph("3. Gunakan Full Source Code Documentation sebagai referensi detail per file.")

        DocStyles.add_horizontal_line(self.doc)
        self.doc.add_page_break()

    def _add_cover_page(self):
        """Add cover page to document."""
        # Add some spacing
        for _ in range(5):
            self.doc.add_paragraph()

        # Main title
        title = self.doc.add_heading(self.options.document_title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = subtitle.add_run(f'\n📚 {self.options.document_subtitle} 📚')
        run.font.size = Pt(16)
        run.font.italic = True


        # Project name
        project_name = os.path.basename(self.root_dir)
        proj = self.doc.add_paragraph()
        proj.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = proj.add_run(f'\nProject: {project_name}')
        run.font.size = Pt(14)
        run.font.bold = True

        # Add spacing
        for _ in range(3):
            self.doc.add_paragraph()

        # Features list
        features = self.doc.add_paragraph()
        features.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = features.add_run(
            '\n\n✅ API Endpoints Registry\n'
            '✅ Database Models Documentation\n'
            '✅ Full Source Code\n'
            '✅ Dependencies List\n'
            '✅ Environment Variables\n'
            '✅ Statistics & Analytics'
        )
        run.font.size = Pt(10)

        # Generation info
        info = self.doc.add_paragraph()
        info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        generation_text = f'\n\nGenerated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        if self.options.organization:
            generation_text += f'\nOrganization: {self.options.organization}'
        if self.options.author_name:
            generation_text += f'\nAuthor: {self.options.author_name}'
        run = info.add_run(generation_text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Stats box on cover (ENHANCED)
        for _ in range(2):
            self.doc.add_paragraph()

        # Create stats table for better visual
        stats_table = self.doc.add_table(rows=1, cols=1)
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = stats_table.rows[0].cells[0]
        DocStyles.set_cell_shading(cell, 'E3F2FD')  # Light blue background

        # Add stats content with better formatting
        stats_para = cell.paragraphs[0]
        stats_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Total Lines - Big and Bold
        run = stats_para.add_run(f'📊 STATISTIK CODEBASE\n\n')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)

        run = stats_para.add_run(f'Total Baris Kode: {self.stats.total_lines:,}\n')
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0)

        run = stats_para.add_run(f'Total File: {self.stats.total_files}\n')
        run.font.size = Pt(12)
        run.font.bold = True

        run = stats_para.add_run(f'Ukuran: {self.stats.total_bytes / 1024:.2f} KB\n\n')
        run.font.size = Pt(10)

        # Language breakdown
        if self.stats.lines_by_type:
            run = stats_para.add_run('BAHASA PEMROGRAMAN:\n')
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204)

            # Sort languages by line count
            sorted_langs = sorted(self.stats.lines_by_type.items(), key=lambda x: x[1], reverse=True)
            top_langs = sorted_langs[:5]  # Top 5 languages

            for ext, lines in top_langs:
                lang_name = Config.LANGUAGE_MAP.get(ext, ext)
                percentage = (lines / self.stats.total_lines * 100) if self.stats.total_lines > 0 else 0

                # Visual bar representation
                bar_length = int(percentage / 5)  # 20 chars max
                bar = '█' * bar_length

                run = stats_para.add_run(f'{lang_name}: {lines:,} baris ({percentage:.1f}%)\n')
                run.font.size = Pt(9)

                run = stats_para.add_run(f'{bar}\n')
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0, 153, 76)

        # Path info
        self.doc.add_paragraph()
        path_info = self.doc.add_paragraph()
        path_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = path_info.add_run(f'Path: {self.root_dir}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)

        # Page break
        self.doc.add_page_break()

    def _add_table_of_contents(self):
        """Add table of contents."""
        self.doc.add_heading('📑 DAFTAR ISI (Table of Contents)', level=1)

        toc_items = [
            "1. Executive Summary",
            "2. Project Overview",
            "3. API Endpoints Registry",
            "4. Database Models",
            "5. Dependencies (requirements.txt)",
            "6. Environment Variables",
            "7. Directory Structure",
            "8. Codebase Statistics",
            "9. Full Source Code Documentation",
            "   9.1. Backend (Python/FastAPI)",
            "   9.2. Frontend (HTML/CSS/JS)",
            "   9.3. Flutter Client",
            "   9.4. Scripts & Configuration",
            "10. Appendix"
        ]

        for item in toc_items:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.left_indent = Inches(0.5)

        DocStyles.add_horizontal_line(self.doc)
        self.doc.add_page_break()

    def _add_project_overview(self):
        """Add project overview section."""
        self.doc.add_heading('2. 📋 PROJECT OVERVIEW', level=1)

        self.doc.add_paragraph(
            "Bagian ini merangkum identitas project, direktori penting, dan cuplikan README "
            "agar pembaca cepat memahami konteks aplikasi."
        )

        overview_table = self.doc.add_table(rows=6, cols=2)
        overview_table.style = 'Table Grid'
        overview_rows = [
            ("Nama Project", os.path.basename(self.root_dir)),
            ("Root Directory", self.root_dir),
            ("Jumlah File", str(self.stats.total_files)),
            ("Jumlah Baris Kode", f"{self.stats.total_lines:,}"),
            ("Waktu Generate", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ("Script Generator", os.path.relpath(__file__, self.root_dir)),
        ]
        for idx, (key, value) in enumerate(overview_rows):
            overview_table.rows[idx].cells[0].text = key
            overview_table.rows[idx].cells[1].text = sanitize_for_xml(value)
            overview_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            DocStyles.set_cell_shading(overview_table.rows[idx].cells[0], "E8E8E8")

        self.doc.add_paragraph()

        # Read README if exists (excerpt for readability)
        readme_path = os.path.join(self.root_dir, 'README.md')
        if os.path.exists(readme_path):
            self.doc.add_heading('README.md (Ringkasan)', level=2)
            try:
                with open(readme_path, 'r', encoding='utf-8') as f:
                    readme_content = f.read()
                lines = readme_content.splitlines()
                excerpt_limit = 120
                excerpt = "\n".join(lines[:excerpt_limit])
                if len(lines) > excerpt_limit:
                    excerpt += (
                        f"\n\n... [dipotong untuk kerapian dokumen, "
                        f"total {len(lines)} baris di README.md]"
                    )
                p = self.doc.add_paragraph(sanitize_for_xml(excerpt))
                p.style = 'Code'
            except Exception as e:
                self.doc.add_paragraph(f"Error reading README: {e}")
        else:
            self.doc.add_paragraph("No README.md found in root directory.")

        DocStyles.add_horizontal_line(self.doc)
        self.doc.add_page_break()

    def _add_api_registry(self):
        """Add API endpoints registry section."""
        self.doc.add_heading('3. 🔌 API ENDPOINTS REGISTRY', level=1)

        self.doc.add_paragraph(
            "Daftar lengkap semua API endpoints yang tersedia dalam aplikasi ini.\n"
            "Endpoints dikelompokkan berdasarkan file dan router."
        )

        # Pre-scan API files
        api_dir = os.path.join(self.root_dir, 'app', 'api')
        if os.path.exists(api_dir):
            all_routes = []

            for filename in os.listdir(api_dir):
                if filename.endswith('.py') and not filename.startswith('__'):
                    filepath = os.path.join(api_dir, filename)

                    # Use regex-based extraction for more reliable results
                    routes = self._extract_routes_regex(filepath)

                    if routes:
                        self.doc.add_heading(f"📄 {filename}", level=2)

                        # Create table for routes
                        table = self.doc.add_table(rows=len(routes) + 1, cols=4)
                        table.style = 'Table Grid'

                        # Header
                        headers = ['Method', 'Path', 'Function', 'Description']
                        for i, header in enumerate(headers):
                            table.rows[0].cells[i].text = header
                            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
                            DocStyles.set_cell_shading(table.rows[0].cells[i], "E8E8E8")

                        # Data
                        for j, route in enumerate(routes, 1):
                            table.rows[j].cells[0].text = route.get('method', '')
                            table.rows[j].cells[1].text = route.get('path', '')
                            table.rows[j].cells[2].text = route.get('function', '')
                            desc = route.get('docstring', '')
                            table.rows[j].cells[3].text = (desc[:60] + '...') if desc and len(desc) > 60 else (desc or '')

                        all_routes.extend(routes)
                        self.doc.add_paragraph()

            self.stats.api_routes = all_routes

            # Summary
            if all_routes:
                self.doc.add_paragraph(f"\n📊 Total API Endpoints: {len(all_routes)}")
        else:
            self.doc.add_paragraph("No API directory found at app/api/")

        DocStyles.add_horizontal_line(self.doc)
        self.doc.add_page_break()

    def _extract_routes_regex(self, filepath):
        """Extract API routes using regex (more reliable than AST for decorators)."""
        routes = []

        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()

            # Get router prefix if available
            prefix_match = re.search(r'router\s*=\s*APIRouter\s*\([^)]*prefix\s*=\s*["\']([^"\']+)["\']', content, re.DOTALL)
            prefix = prefix_match.group(1) if prefix_match else ""

            # Pattern 1: Match @router.method("path") with various formats
            # Handles multiline decorators and keyword arguments
            pattern1 = r'@(?:router|app)\.(get|post|put|delete|patch)\s*\(\s*["\']([^"\']+)["\']'

            # Find all decorator matches
            decorator_matches = list(re.finditer(pattern1, content, re.IGNORECASE))

            for dec_match in decorator_matches:
                method = dec_match.group(1).upper()
                path = dec_match.group(2)

                # Find the function definition after this decorator
                # Start searching from decorator position
                search_start = dec_match.end()
                func_pattern = r'(?:async\s+)?def\s+(\w+)\s*\('
                func_match = re.search(func_pattern, content[search_start:search_start+500])

                if func_match:
                    func_name = func_match.group(1)

                    # Try to get docstring
                    docstring = ""
                    docstring_start = search_start + func_match.end()
                    docstring_pattern = r'^\s*["\'\]{3}([^"\']+)["\'\]{3}'
                    doc_content = content[docstring_start:docstring_start+500]
                    doc_match = re.search(r'"""([^"]+)"""', doc_content)
                    if doc_match:
                        docstring = doc_match.group(1).strip().split('\n')[0]

                    routes.append({
                        'method': method,
                        'path': prefix + path,
                        'function': func_name,
                        'docstring': docstring
                    })

        except Exception as e:
            print(f"  Warning: Could not parse {filepath}: {e}")

        return routes

    def _add_database_models(self):
        """Add database models documentation section."""
        self.doc.add_heading('4. 🗃️ DATABASE MODELS', level=1)

        self.doc.add_paragraph(
            "Dokumentasi semua model database (SQLAlchemy) yang digunakan dalam aplikasi."
        )

        # Scan models directory
        models_dir = os.path.join(self.root_dir, 'app', 'models')
        if os.path.exists(models_dir):
            all_models = []

            for filename in os.listdir(models_dir):
                if filename.endswith('.py') and not filename.startswith('__'):
                    filepath = os.path.join(models_dir, filename)
                    analysis = PythonAnalyzer.analyze_file(filepath)

                    if 'error' not in analysis:
                        # Check for classes that look like models
                        for cls in analysis.get('classes', []):
                            if any(base in ['Base', 'Model', 'db.Model'] for base in cls.get('bases', [])):
                                all_models.append({
                                    'file': filename,
                                    'class': cls
                                })

                                self.doc.add_heading(f"📊 {cls['name']}", level=2)
                                self.doc.add_paragraph(f"File: {filename}")

                                if cls.get('docstring'):
                                    self.doc.add_paragraph(f"Description: {cls['docstring']}")

                                # List methods/properties
                                if cls.get('methods'):
                                    self.doc.add_heading("Properties & Methods:", level=3)
                                    for method in cls['methods']:
                                        p = self.doc.add_paragraph()
                                        run = p.add_run(f"• {method['name']}()")
                                        if method.get('docstring'):
                                            p.add_run(f" - {method['docstring'][:80]}")

                                self.doc.add_paragraph()

            self.stats.models = all_models
        else:
            self.doc.add_paragraph("No models directory found at app/models/")

        DocStyles.add_horizontal_line(self.doc)
        self.doc.add_page_break()

    def _add_dependencies_section(self):
        """Add dependencies section."""
        self.doc.add_heading('5. 📦 DEPENDENCIES', level=1)

        if self.stats.dependencies:
            self.doc.add_paragraph(
                f"Total {len(self.stats.dependencies)} dependencies from requirements.txt:"
            )

            # Create table
            table = self.doc.add_table(rows=len(self.stats.dependencies) + 1, cols=2)
            table.style = 'Table Grid'

            # Header
            table.rows[0].cells[0].text = "Package"
            table.rows[0].cells[1].text = "Version"
            table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
            DocStyles.set_cell_shading(table.rows[0].cells[0], "E8E8E8")
            DocStyles.set_cell_shading(table.rows[0].cells[1], "E8E8E8")

            # Data
            for i, dep in enumerate(self.stats.dependencies, 1):
                table.rows[i].cells[0].text = dep.get('name', '')
                table.rows[i].cells[1].text = dep.get('version', 'latest')
        else:
            self.doc.add_paragraph("No requirements.txt found.")

        DocStyles.add_horizontal_line(self.doc)
        self.doc.add_page_break()

    def _add_environment_variables(self):
        """Add environment variables section."""
        self.doc.add_heading('6. ⚙️ ENVIRONMENT VARIABLES', level=1)

        if self.stats.env_vars:
            self.doc.add_paragraph(
                f"Total {len(self.stats.env_vars)} environment variables from .env.example:"
            )

            # Create table
            table = self.doc.add_table(rows=len(self.stats.env_vars) + 1, cols=2)
            table.style = 'Table Grid'

            # Header
            table.rows[0].cells[0].text = "Variable Name"
            table.rows[0].cells[1].text = "Example Value"
            table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
            DocStyles.set_cell_shading(table.rows[0].cells[0], "E8E8E8")
            DocStyles.set_cell_shading(table.rows[0].cells[1], "E8E8E8")

            # Data
            for i, env in enumerate(self.stats.env_vars, 1):
                table.rows[i].cells[0].text = env.get('name', '')
                # Mask sensitive values
                value = env.get('example_value', '')
                if any(x in env.get('name', '').upper() for x in ['PASSWORD', 'SECRET', 'KEY', 'TOKEN']):
                    value = '********' if value else '(not set)'
                table.rows[i].cells[1].text = value
        else:
            self.doc.add_paragraph("No .env.example found.")

        DocStyles.add_horizontal_line(self.doc)
        self.doc.add_page_break()

    def _add_directory_structure(self):
        """Add directory structure section."""
        self.doc.add_heading('7. 📁 DIRECTORY STRUCTURE', level=1)

        self.doc.add_paragraph(
            "Berikut adalah struktur direktori lengkap dari project ini:"
        )

        # Generate tree
        tree_lines = DirectoryTree.generate(self.root_dir, max_depth=self.options.max_tree_depth)

        # Add tree to document
        for line in tree_lines:
            p = self.doc.add_paragraph(sanitize_for_xml(line))
            p.style = 'Code'
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        DocStyles.add_horizontal_line(self.doc)
        self.doc.add_page_break()

    def _add_statistics_section(self):
        """Add statistics section."""
        self.doc.add_heading('8. 📊 CODEBASE STATISTICS', level=1)

        # Statistics summary
        stats = self.stats.get_summary()

        summary_table = self.doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'

        summary_data = [
            ("Total Files", str(stats['total_files'])),
            ("Total Lines of Code", f"{stats['total_lines']:,}"),
            ("Total Size", f"{stats['total_bytes'] / 1024:.2f} KB"),
            ("API Endpoints", str(len(stats['api_routes']))),
            ("Database Models", str(len(stats['models']))),
            ("Scan Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]

        for i, (key, value) in enumerate(summary_data):
            cell_k = summary_table.rows[i].cells[0]
            cell_k.text = key
            cell_k.paragraphs[0].runs[0].font.bold = True
            DocStyles.set_cell_shading(cell_k, "E8E8E8")
            summary_table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        # Language breakdown
        self.doc.add_heading("Language Breakdown:", level=2)
        lang_table = self.doc.add_table(rows=len(stats['files_by_type']) + 1, cols=3)
        lang_table.style = 'Table Grid'

        headers = ['Extension', 'Files', 'Lines']
        for i, h in enumerate(headers):
            lang_table.rows[0].cells[i].text = h
            lang_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            DocStyles.set_cell_shading(lang_table.rows[0].cells[i], "F0F0F0")

        for i, (ext, count) in enumerate(sorted(stats['files_by_type'].items()), 1):
            lang_table.rows[i].cells[0].text = ext or "(no ext)"
            lang_table.rows[i].cells[1].text = str(count)
            lang_table.rows[i].cells[2].text = str(stats['lines_by_type'].get(ext, 0))

        self.doc.add_page_break()

    def _add_all_files_documentation(self):
        """Add documentation for all files."""
        self.doc.add_heading('9. 📝 FULL SOURCE CODE DOCUMENTATION', level=1)

        if self.options.include_source_code:
            self.doc.add_paragraph(
                "Dokumentasi lengkap dari seluruh source code dalam project ini. "
                "Setiap file ditampilkan dengan path lengkap, informasi bahasa pemrograman, "
                "serta konten source code yang dapat di-review secara mendalam."
            )
        else:
            self.doc.add_paragraph(
                "Dokumentasi ringkas per file (tanpa konten source code lengkap). "
                "Mode ini cocok untuk laporan manajerial agar dokumen tetap ringan."
            )

        DocStyles.add_horizontal_line(self.doc)

        # Process all files
        current_section = None

        for dirpath, dirnames, filenames in os.walk(self.root_dir):
            # Skip excluded directories
            dirnames[:] = [d for d in dirnames if d not in Config.EXCLUDE_DIRS]

            for filename in sorted(filenames):
                filepath = os.path.join(dirpath, filename)
                rel_path = os.path.relpath(filepath, self.root_dir)

                # Get extension
                _, ext = os.path.splitext(filename)
                ext = ext.lower()

                # Skip excluded files
                if ext in Config.EXCLUDE_EXTENSIONS or filename in Config.EXCLUDE_FILES:
                    continue

                # Skip the output file itself
                if filename == self.output_file:
                    continue

                # Determine section
                section = self._determine_section(rel_path)
                if section != current_section:
                    current_section = section
                    self.doc.add_heading(f"9.x. {section}", level=2)

                # Process file
                self._add_file_documentation(filepath, rel_path, ext)

    def _determine_section(self, rel_path):
        """Determine which section a file belongs to."""
        rel_path_lower = rel_path.lower()

        if rel_path_lower.startswith('app'):
            return "🐍 Backend (Python/FastAPI)"
        elif rel_path_lower.startswith('templates') or rel_path_lower.startswith('static'):
            return "🌐 Frontend (Templates/Static)"
        elif rel_path_lower.startswith('flutter_client'):
            return "🎯 Flutter Client"
        elif rel_path_lower.startswith('scripts'):
            return "⚙️ Scripts"
        else:
            return "📦 Configuration & Other Files"

    def _add_file_documentation(self, filepath, rel_path, ext):
        """Add documentation for a single file."""
        print(f"📄 Processing: {rel_path}")

        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
        except Exception as e:
            self.doc.add_paragraph(f"Error reading {rel_path}: {e}")
            return

        # Update statistics (REMOVED: Done in pre-scan)
        # self.stats.add_file(rel_path, content, ext)

        # File header
        self.doc.add_heading(f"📄 {rel_path}", level=3)

        # File metadata table
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        # Metadata
        metadata = [
            ("Path", rel_path),
            ("Language", Config.LANGUAGE_MAP.get(ext, f"Unknown ({ext})")),
            ("Lines", str(len(content.splitlines()))),
            ("Size", f"{len(content.encode('utf-8'))} bytes")
        ]

        for i, (key, value) in enumerate(metadata):
            cell_key = table.rows[i].cells[0]
            cell_key.text = key
            cell_key.paragraphs[0].runs[0].font.bold = True
            DocStyles.set_cell_shading(cell_key, "E8E8E8")

            cell_value = table.rows[i].cells[1]
            cell_value.text = value

        self.doc.add_paragraph()  # Spacing

        # For Python files, add analysis
        if ext == '.py':
            analysis = PythonAnalyzer.analyze_file(filepath)
            if 'error' not in analysis:
                self._add_python_analysis(analysis)

        # For JavaScript files, add analysis
        elif ext == '.js':
            analysis = JavaScriptAnalyzer.analyze_file(filepath)
            if 'error' not in analysis:
                self._add_javascript_analysis(analysis)

        # For HTML files, add analysis
        elif ext == '.html':
            analysis = HTMLAnalyzer.analyze_file(filepath)
            if 'error' not in analysis:
                self._add_html_analysis(analysis)

        # For Dart files, add analysis
        elif ext == '.dart':
            analysis = DartAnalyzer.analyze_file(filepath)
            if 'error' not in analysis:
                self._add_dart_analysis(analysis)

        if self.options.include_source_code:
            # Add source code
            self.doc.add_heading("Source Code:", level=4)

            if not content.strip():
                self.doc.add_paragraph("(File is empty)")
            else:
                lines = content.splitlines()
                total_lines = len(lines)
                display_lines = lines
                truncated = False

                max_lines = max(1, int(self.options.max_code_lines_per_file))
                if total_lines > max_lines:
                    truncated = True
                    head_count = max_lines // 2
                    tail_count = max_lines - head_count
                    display_lines = (
                        lines[:head_count]
                        + [f"... [TRUNCATED: {total_lines - max_lines} lines omitted for readability] ..."]
                        + lines[-tail_count:]
                    )

                # Add line numbers
                numbered_content = []
                for i, line in enumerate(display_lines, 1):
                    numbered_content.append(f"{i:4d} | {line}")

                code_text = sanitize_for_xml("\n".join(numbered_content))

                p = self.doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)

                if truncated:
                    note = self.doc.add_paragraph(
                        f"Catatan: File dipotong untuk kerapian dokumen. "
                        f"Total asli {total_lines} baris, ditampilkan {len(display_lines)} baris."
                    )
                    if note.runs:
                        note.runs[0].font.italic = True
        else:
            self.doc.add_paragraph(
                "Source code lengkap tidak disertakan (mode ringkas aktif)."
            )
            self.doc.add_paragraph(
                "Gunakan opsi include-source untuk memasukkan source code per file."
            )

        DocStyles.add_horizontal_line(self.doc)

    def _add_python_analysis(self, analysis):
        """Add Python file analysis section."""
        if analysis.get('module_docstring'):
            self.doc.add_heading("Module Docstring:", level=4)
            p = self.doc.add_paragraph(sanitize_for_xml(analysis['module_docstring']))
            p.style = 'Code'

        if analysis.get('classes'):
            self.doc.add_heading("Classes:", level=4)
            for cls in analysis['classes']:
                p = self.doc.add_paragraph()
                run = p.add_run(f"• class {cls['name']} (line {cls['line']})")
                run.font.bold = True

                if cls.get('docstring'):
                    docstr = sanitize_for_xml(cls['docstring'][:200])
                    self.doc.add_paragraph(f"  {docstr}...")

                for method in cls.get('methods', []):
                    args = ', '.join(method['args'])
                    self.doc.add_paragraph(f"    └─ def {method['name']}({args})")

        if analysis.get('functions'):
            self.doc.add_heading("Functions:", level=4)
            for func in analysis['functions']:
                args = ', '.join(func['args'])
                p = self.doc.add_paragraph()
                run = p.add_run(f"• def {func['name']}({args}) - line {func['line']}")
                run.font.bold = True

                if func.get('docstring'):
                    docstr = sanitize_for_xml(func['docstring'][:150])
                    self.doc.add_paragraph(f"  {docstr}...")

    def _add_javascript_analysis(self, analysis):
        """Add JavaScript file analysis section."""
        if analysis.get('classes'):
            self.doc.add_heading("Classes:", level=4)
            for cls in analysis['classes']:
                extends = f" extends {cls['extends']}" if cls.get('extends') else ""
                self.doc.add_paragraph(f"• class {cls['name']}{extends}")

        if analysis.get('functions'):
            self.doc.add_heading("Functions:", level=4)
            for func in analysis['functions']:
                func_type = "arrow" if func.get('type') == 'arrow' else "function"
                self.doc.add_paragraph(f"• {func['name']} ({func_type})")

    def _add_html_analysis(self, analysis):
        """Add HTML file analysis section."""
        if analysis.get('title'):
            self.doc.add_paragraph(f"Page Title: {analysis['title']}")

        if analysis.get('scripts'):
            self.doc.add_heading("Scripts:", level=4)
            for script in analysis['scripts'][:5]:  # Limit to 5
                self.doc.add_paragraph(f"• {script}")

        if analysis.get('forms'):
            self.doc.add_heading("Forms:", level=4)

    def _add_dart_analysis(self, analysis):
        """Add Dart file analysis section."""
        if analysis.get('imports'):
            self.doc.add_heading("Imports:", level=4)
            # Show top 5 imports
            for imp in analysis['imports'][:5]:
                self.doc.add_paragraph(f"• {imp}")
            if len(analysis['imports']) > 5:
                self.doc.add_paragraph(f"  ... and {len(analysis['imports']) - 5} more")

        if analysis.get('classes'):
            self.doc.add_heading("Classes:", level=4)
            for cls in analysis['classes']:
                text = f"• class {cls['name']}"
                if cls.get('extends'):
                    text += f" extends {cls['extends']}"
                if cls.get('implements'):
                    text += f" implements {cls['implements']}"
                self.doc.add_paragraph(text)

        if analysis.get('functions'):
            self.doc.add_heading("Functions detected:", level=4)
            # Filter duplicates and common methods
            funcs = sorted(list(set(analysis['functions'])))
            for func in funcs:
                if func not in ['build', 'initState', 'dispose']: # Skip common widget methods to reduce noise
                   self.doc.add_paragraph(f"• {func}()")

    def _add_appendix(self):
        """Add appendix with additional information."""
        self.doc.add_page_break()
        self.doc.add_heading('10. 📎 APPENDIX', level=1)

        # Statistics summary
        stats = self.stats.get_summary()

        self.doc.add_heading("A. Complete Statistics Summary", level=2)

        # Summary table
        summary_table = self.doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'

        summary_data = [
            ("Total Files", str(stats['total_files'])),
            ("Total Lines of Code", f"{stats['total_lines']:,}"),
            ("Total Size", f"{stats['total_bytes'] / 1024:.2f} KB"),
            ("API Endpoints", str(len(stats['api_routes']))),
            ("Database Models", str(len(stats['models']))),
            ("Generation Date", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]

        for i, (key, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = key
            summary_table.rows[i].cells[1].text = value

        self.doc.add_paragraph()

        # Files by type
        self.doc.add_heading("B. Files by Type", level=2)

        type_table = self.doc.add_table(rows=len(stats['files_by_type']) + 1, cols=3)
        type_table.style = 'Table Grid'

        # Header
        headers = ['Extension', 'File Count', 'Lines']
        for i, header in enumerate(headers):
            type_table.rows[0].cells[i].text = header
            type_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

        # Data
        for i, (ext, count) in enumerate(sorted(stats['files_by_type'].items()), 1):
            type_table.rows[i].cells[0].text = ext or "(no extension)"
            type_table.rows[i].cells[1].text = str(count)
            type_table.rows[i].cells[2].text = str(stats['lines_by_type'].get(ext, 0))

        self.doc.add_paragraph()

        # Largest files
        self.doc.add_heading("C. Top 10 Largest Files (by lines)", level=2)

        if stats['largest_files']:
            large_table = self.doc.add_table(rows=len(stats['largest_files'][:10]) + 1, cols=3)
            large_table.style = 'Table Grid'

            headers = ['File Path', 'Lines', 'Size']
            for i, header in enumerate(headers):
                large_table.rows[0].cells[i].text = header
                large_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

            for i, file_info in enumerate(stats['largest_files'][:10], 1):
                large_table.rows[i].cells[0].text = file_info['path']
                large_table.rows[i].cells[1].text = str(file_info['lines'])
                large_table.rows[i].cells[2].text = f"{file_info['size'] / 1024:.2f} KB"

        # Final note
        self.doc.add_paragraph()
        DocStyles.add_horizontal_line(self.doc)

        final_note = self.doc.add_paragraph()
        final_note.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = final_note.add_run(
            "\n\n— End of Documentation —\n"
            "Generated by Full Codebase Documentation Generator v4.0\n"
            "100% Complete Application Documentation"
        )
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)


# ============================================================================
# ENTRY POINT
# ============================================================================
def main():
    """Main entry point."""
    default_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    parser = argparse.ArgumentParser(
        description="Generate full DOCX documentation for a codebase.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "root_dir",
        nargs="?",
        default=default_root,
        help="Root directory project yang akan didokumentasikan",
    )
    parser.add_argument(
        "output_file",
        nargs="?",
        default="FULL_CODEBASE_DOCUMENTATION.docx",
        help="Nama file output DOCX",
    )
    parser.add_argument("--title", default=GeneratorOptions.document_title, help="Judul dokumen")
    parser.add_argument("--subtitle", default=GeneratorOptions.document_subtitle, help="Subjudul dokumen")
    parser.add_argument("--author", default=GeneratorOptions.author_name, help="Nama author dokumen")
    parser.add_argument("--organization", default="", help="Nama organisasi/institusi")
    parser.add_argument(
        "--no-source",
        action="store_true",
        help="Jangan sertakan source code lengkap (mode ringkas)",
    )
    parser.add_argument(
        "--max-code-lines",
        type=int,
        default=GeneratorOptions.max_code_lines_per_file,
        help="Batas baris code per file pada mode include source",
    )
    parser.add_argument(
        "--max-tree-depth",
        type=int,
        default=GeneratorOptions.max_tree_depth,
        help="Kedalaman maksimum directory tree",
    )
    parser.add_argument(
        "--auto-install-deps",
        action="store_true",
        help="Auto install python-docx jika belum tersedia",
    )

    args = parser.parse_args()
    load_docx_symbols(auto_install=args.auto_install_deps)

    options = GeneratorOptions(
        document_title=args.title,
        document_subtitle=args.subtitle,
        author_name=args.author,
        organization=args.organization,
        include_source_code=not args.no_source,
        max_code_lines_per_file=max(1, args.max_code_lines),
        max_tree_depth=max(1, args.max_tree_depth),
    )

    # Generate documentation
    generator = FullDocumentationGenerator(args.root_dir, args.output_file, options=options)
    output_path = generator.generate()

    return output_path


if __name__ == "__main__":
    main()
