from io import BytesIO
import re
from typing import Dict, Any, List

from docx import Document
from docx.oxml.ns import qn

from app.core.assessment_docx_generator import (
    PAN_TEMPLATE_PATH,
    PAP_TEMPLATE_PATH,
    generate_assessment_docx,
)


UNRESOLVED_PATTERN = re.compile(r"\[[^\]]+\]|_{3,}")


def _collect_doc_text(document: Document) -> str:
    chunks: List[str] = []
    chunks.extend((p.text or "") for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text or "")
    return "\n".join(chunks)


def _cell_fill(cell: Any) -> str:
    tc_pr = getattr(cell._tc, "tcPr", None)
    if tc_pr is None:
        return ""
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return ""
    return str(shd.get(qn("w:fill")) or "")


def _classify_pan(score: float) -> str:
    if score >= 90:
        return "A"
    if score >= 80:
        return "B"
    if score >= 70:
        return "C"
    if score >= 60:
        return "D"
    return "E"


def _classify_pap(score: float) -> str:
    if score >= 90:
        return "A"
    if score >= 80:
        return "B"
    if score >= 70:
        return "C"
    if score >= 60:
        return "D"
    return "E"


def _build_payload(participant_count: int = 35) -> Dict[str, Any]:
    um_map = {
        "A": {"range": "90 - 95", "predicate": "Sangat Baik", "min": 90, "max": 95},
        "B": {"range": "85 - 89", "predicate": "Baik", "min": 85, "max": 89},
        "C": {"range": "80 - 84", "predicate": "Cukup", "min": 80, "max": 84},
        "D": {"range": "75 - 79", "predicate": "Kurang", "min": 75, "max": 79},
        "E": {"range": "70 - 74", "predicate": "Sangat Kurang", "min": 70, "max": 74},
    }
    participants = []
    scores = []
    for index in range(participant_count):
        score = float(max(45, 96 - index))
        scores.append(score)
        pan_letter = _classify_pan(score)
        pap_letter = _classify_pap(score)
        um_cfg = um_map[pan_letter]
        um_score = int(round((um_cfg["min"] + um_cfg["max"]) / 2))
        participants.append(
            {
                "rank": index + 1,
                "name": f"Peserta {index + 1}",
                "student_class": "XII C",
                "score": score,
                "pan_letter": pan_letter,
                "pan_scale10": max(0, min(10, int(round(score / 10.0)))),
                "t_score": 50.0 + ((score - 75.0) / 10.0) * 10.0,
                "um_score": um_score,
                "um_category": um_cfg["range"],
                "um_predicate": um_cfg["predicate"],
                "pap_letter": pap_letter,
                "pap_status": "TUNTAS" if score >= 78.0 else "TIDAK TUNTAS",
            }
        )

    letter_counts = {grade: 0 for grade in ["A", "B", "C", "D", "E"]}
    for participant in participants:
        letter_counts[participant["pan_letter"]] += 1
    total = float(participant_count)

    pan_distribution = []
    for grade in ["A", "B", "C", "D", "E"]:
        count = letter_counts[grade]
        pan_distribution.append(
            {
                "grade": grade,
                "range": "-",
                "count": count,
                "percentage": round((count / total) * 100.0, 2) if total else 0.0,
            }
        )

    um_summary = []
    for grade in ["A", "B", "C", "D", "E"]:
        filtered = [p for p in participants if p["pan_letter"] == grade]
        um_summary.append(
            {
                "category": grade,
                "pan_range": "-",
                "count": len(filtered),
                "percentage": round((len(filtered) / total) * 100.0, 2) if total else 0.0,
                "um_range": "-",
                "predicate": "-",
                "student_names": ", ".join(p["name"] for p in filtered[:8]) if filtered else "-",
            }
        )

    pass_count = sum(1 for p in participants if p["pap_status"] == "TUNTAS")
    fail_count = participant_count - pass_count
    pap_counts = {grade: 0 for grade in ["A", "B", "C", "D", "E"]}
    for participant in participants:
        pap_counts[participant["pap_letter"]] += 1
    pap_distribution = [
        {
            "grade": "A",
            "range": "90 - 100",
            "count": pap_counts["A"],
            "percentage": round((pap_counts["A"] / total) * 100.0, 2) if total else 0.0,
        },
        {
            "grade": "B",
            "range": "80 - 89",
            "count": pap_counts["B"],
            "percentage": round((pap_counts["B"] / total) * 100.0, 2) if total else 0.0,
        },
        {
            "grade": "C",
            "range": "70 - 79",
            "count": pap_counts["C"],
            "percentage": round((pap_counts["C"] / total) * 100.0, 2) if total else 0.0,
        },
        {
            "grade": "D",
            "range": "60 - 69",
            "count": pap_counts["D"],
            "percentage": round((pap_counts["D"] / total) * 100.0, 2) if total else 0.0,
        },
        {
            "grade": "E",
            "range": "< 60",
            "count": pap_counts["E"],
            "percentage": round((pap_counts["E"] / total) * 100.0, 2) if total else 0.0,
        },
    ]

    return {
        "exam": {
            "title": "ASA Bahasa Inggris Kelas XII",
            "date_text": "19 April 2026 08:00 WIB",
            "teacher_name": "Guru Bahasa Inggris",
            "subject": "Bahasa Inggris",
        },
        "class_name": "XII C",
        "generated_at": "19 April 2026 14:30 WIB",
        "stats": {
            "participant_count": participant_count,
            "average": round(sum(scores) / len(scores), 2),
            "std_dev": 10.0,
            "highest": max(scores),
            "lowest": min(scores),
        },
        "pan": {
            "mean": round(sum(scores) / len(scores), 2),
            "std_dev": 10.0,
            "score_range": max(scores) - min(scores),
            "class_count": 6,
            "interval": round((max(scores) - min(scores)) / 6.0, 2),
            "thresholds": {
                "a_min": 90.0,
                "b_min": 80.0,
                "c_min": 70.0,
                "d_min": 60.0,
            },
            "scale10_thresholds": {
                "10_min": 97.5,
                "9_min": 92.5,
                "8_min": 87.5,
                "7_min": 82.5,
                "6_min": 77.5,
                "5_min": 72.5,
                "4_min": 67.5,
                "3_min": 62.5,
                "2_min": 57.5,
                "1_min": 52.5,
                "0_max": 52.5,
            },
            "letter_distribution": pan_distribution,
            "um_conversion_summary": um_summary,
        },
        "pap": {
            "kkm": 78.0,
            "smi": 100.0,
            "pass_count": pass_count,
            "fail_count": fail_count,
            "pass_percentage": round((pass_count / total) * 100.0, 2) if total else 0.0,
            "grade_distribution": pap_distribution,
        },
        "participants": participants,
    }


def test_generate_pan_docx_fills_template_without_unresolved_placeholders() -> None:
    payload = _build_payload(participant_count=35)
    output_bytes = generate_assessment_docx("pan", payload)

    output_doc = Document(BytesIO(output_bytes))
    template_doc = Document(str(PAN_TEMPLATE_PATH))
    output_text = _collect_doc_text(output_doc)

    assert len(output_doc.tables) == len(template_doc.tables) + 1
    assert "Lampiran Data Otomatis PAN" not in output_text
    assert not UNRESOLVED_PATTERN.search(output_text)
    assert len(output_doc.tables[6].rows) >= 36  # header + 35 peserta
    assert "Nilai UM" in output_text
    assert "Predikat" in output_text
    assert "Peserta 35" in output_text
    assert "Mengetahui," in output_text
    assert "Kepala MAN 1 Rokan Hulu" in output_text
    assert "Rokan Hulu, 19 April 2026" in output_text
    assert "Guru Mata Pelajaran Bahasa Inggris" in output_text
    assert output_text.count("NIP.") >= 2
    assert "Guru Pelaksana," not in output_text
    assert "[Nama guru pelaksana]" not in output_text

    pan_grade_table = output_doc.tables[2]
    pan_grade_headers = [cell.text for cell in pan_grade_table.rows[0].cells]
    assert len(pan_grade_headers) == 3
    assert pan_grade_headers == ["Kategori", "Rumus Batas Bawah", "Batas Skor"]
    assert "Keterangan" not in pan_grade_headers

    participant_table = output_doc.tables[6]
    header_cells = participant_table.rows[0].cells
    row1_cells = participant_table.rows[1].cells
    header_fill = _cell_fill(header_cells[0])
    row_fill = _cell_fill(row1_cells[0])
    assert header_fill != ""
    assert row_fill != ""
    assert len(header_cells) == 10
    for idx in [8, 9]:
        assert _cell_fill(header_cells[idx]) == header_fill
        assert _cell_fill(row1_cells[idx]) == row_fill


def test_generate_pap_docx_fills_template_without_unresolved_placeholders() -> None:
    payload = _build_payload(participant_count=35)
    output_bytes = generate_assessment_docx("pap", payload)

    output_doc = Document(BytesIO(output_bytes))
    template_doc = Document(str(PAP_TEMPLATE_PATH))
    output_text = _collect_doc_text(output_doc)

    assert len(output_doc.tables) == len(template_doc.tables)
    assert "Lampiran Data Otomatis PAP" not in output_text
    assert not UNRESOLVED_PATTERN.search(output_text)
    assert len(output_doc.tables[4].rows) >= 36  # header + 35 peserta
    assert "Peserta 35" in output_text
